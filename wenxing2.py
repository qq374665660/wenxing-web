import openpyxl
from collections import defaultdict, Counter
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os
import logging
import tempfile
import re
import math
from docx.shared import Pt, RGBColor, Inches, Cm  # <--- 加上 Cm
# ====================== 必须补全的 import（修复 NameError） ======================
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
# ==============================================================================
# 设置日志
logging.basicConfig(filename='foundation_analysis.log', level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s', encoding='utf-8')

# --- 全局变量和数据表 ---

# 硬编码平均附加应力系数表
ALPHA_TABLE = [
    ['z/b', 1, 1.2, 1.4, 1.6, 1.8, 2, 2.4, 2.8, 3.2, 3.6, 4, 5, 10],
    [0, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25, 0.25],
    [0.2, 0.2496, 0.2497, 0.2497, 0.2498, 0.2498, 0.2498, 0.2498, 0.2498, 0.2498, 0.2498, 0.2498, 0.2498, 0.2498],
    [0.4, 0.2474, 0.2479, 0.2481, 0.2483, 0.2483, 0.2484, 0.2485, 0.2485, 0.2485, 0.2485, 0.2485, 0.2485, 0.2485],
    [0.6, 0.2423, 0.2437, 0.2444, 0.2449, 0.2451, 0.2452, 0.2454, 0.2455, 0.2455, 0.2455, 0.2455, 0.2455, 0.2456],
    [0.8, 0.2346, 0.2372, 0.2387, 0.2395, 0.24, 0.2403, 0.2407, 0.2408, 0.2409, 0.2409, 0.241, 0.241, 0.241],
    [1, 0.2252, 0.2291, 0.2313, 0.2326, 0.2335, 0.234, 0.2346, 0.2349, 0.2351, 0.2352, 0.2352, 0.2353, 0.2353],
    [1.2, 0.2149, 0.2199, 0.2229, 0.2248, 0.226, 0.2268, 0.2278, 0.2282, 0.2285, 0.2286, 0.2287, 0.2288, 0.2289],
    [1.4, 0.2043, 0.2102, 0.214, 0.2146, 0.218, 0.2191, 0.2204, 0.2211, 0.2215, 0.2217, 0.2218, 0.222, 0.2221],
    [1.6, 0.1939, 0.2006, 0.2049, 0.2079, 0.2099, 0.2113, 0.213, 0.2138, 0.2143, 0.2146, 0.2148, 0.215, 0.2152],
    [1.8, 0.184, 0.1912, 0.196, 0.1994, 0.2018, 0.2034, 0.2055, 0.2066, 0.2073, 0.2077, 0.2079, 0.2082, 0.2084],
    [2, 0.1746, 0.1822, 0.1875, 0.1912, 0.198, 0.1958, 0.1982, 0.1996, 0.2004, 0.2009, 0.2012, 0.2015, 0.2018],
    [2.2, 0.1659, 0.1737, 0.1793, 0.1883, 0.1862, 0.1883, 0.1911, 0.1927, 0.1937, 0.1943, 0.1947, 0.1952, 0.1955],
    [2.4, 0.1578, 0.1657, 0.1715, 0.1757, 0.1789, 0.1812, 0.1843, 0.1862, 0.1873, 0.1873, 0.1885, 0.189, 0.1895],
    [2.6, 0.1503, 0.1583, 0.1642, 0.1686, 0.1719, 0.1745, 0.1779, 0.1799, 0.1812, 0.1812, 0.1825, 0.1832, 0.1838],
    [2.8, 0.1433, 0.1514, 0.1574, 0.1619, 0.1654, 0.168, 0.1717, 0.1739, 0.1753, 0.1753, 0.1769, 0.1777, 0.1784],
    [3, 0.1369, 0.1449, 0.151, 0.1556, 0.1592, 0.1619, 0.1658, 0.1682, 0.1698, 0.1708, 0.1715, 0.1725, 0.1733],
    [3.2, 0.131, 0.139, 0.145, 0.1497, 0.1533, 0.1562, 0.1602, 0.1628, 0.1645, 0.1657, 0.1664, 0.1675, 0.1685],
    [3.4, 0.1256, 0.1334, 0.1394, 0.1441, 0.1478, 0.1508, 0.155, 0.1577, 0.1595, 0.1607, 0.1616, 0.1628, 0.1639],
    [3.6, 0.1205, 0.1282, 0.1342, 0.1389, 0.1427, 0.1456, 0.15, 0.1528, 0.1548, 0.1561, 0.157, 0.1583, 0.1595],
    [3.8, 0.1158, 0.1234, 0.1293, 0.134, 0.1378, 0.1408, 0.1452, 0.1482, 0.1502, 0.1516, 0.1526, 0.1541, 0.1554],
    [4, 0.1114, 0.1189, 0.1248, 0.1294, 0.1332, 0.1362, 0.1408, 0.1438, 0.1459, 0.1474, 0.1485, 0.15, 0.1516],
    [4.2, 0.1073, 0.1147, 0.1205, 0.1251, 0.1289, 0.1319, 0.1365, 0.1396, 0.1418, 0.1434, 0.1445, 0.1462, 0.1479],
    [4.4, 0.1035, 0.1107, 0.1164, 0.121, 0.1248, 0.1279, 0.1325, 0.1357, 0.1379, 0.1396, 0.1407, 0.1425, 0.1444],
    [4.6, 0.1, 0.107, 0.1127, 0.1172, 0.1209, 0.124, 0.1287, 0.1319, 0.1342, 0.1359, 0.1371, 0.139, 0.141],
    [4.8, 0.0967, 0.1036, 0.1091, 0.1136, 0.1173, 0.1204, 0.125, 0.1283, 0.1307, 0.1324, 0.1337, 0.1357, 0.1379],
    [5, 0.0935, 0.1003, 0.1057, 0.1102, 0.1139, 0.1169, 0.1216, 0.1249, 0.1273, 0.1291, 0.1304, 0.1325, 0.1348],
    [5.2, 0.0906, 0.0972, 0.1026, 0.107, 0.1106, 0.1136, 0.1183, 0.1217, 0.1241, 0.1259, 0.1273, 0.1295, 0.132],
    [5.4, 0.0878, 0.0943, 0.0996, 0.1039, 0.1075, 0.1105, 0.1152, 0.1186, 0.121, 0.1229, 0.1243, 0.1265, 0.1292],
    [5.6, 0.0852, 0.0916, 0.0968, 0.101, 0.1046, 0.1076, 0.1122, 0.1156, 0.1181, 0.12, 0.1215, 0.1238, 0.1266],
    [5.8, 0.0828, 0.089, 0.0941, 0.0983, 0.1018, 0.1047, 0.1094, 0.1128, 0.1153, 0.1172, 0.1187, 0.1211, 0.124],
    [6, 0.0805, 0.0866, 0.0916, 0.0957, 0.0991, 0.1021, 0.1067, 0.1101, 0.1126, 0.1146, 0.1161, 0.1186, 0.1216],
    [6.2, 0.0783, 0.0842, 0.0891, 0.0932, 0.0966, 0.0995, 0.1041, 0.1075, 0.1101, 0.112, 0.1136, 0.1161, 0.1193],
    [6.4, 0.0762, 0.082, 0.0869, 0.0909, 0.0942, 0.0971, 0.1016, 0.105, 0.1076, 0.1096, 0.1111, 0.1137, 0.1171],
    [6.6, 0.0742, 0.0799, 0.0847, 0.0886, 0.0919, 0.0948, 0.0993, 0.1027, 0.1053, 0.1073, 0.1088, 0.1114, 0.1149],
    [6.8, 0.0723, 0.0779, 0.0826, 0.0865, 0.0898, 0.0926, 0.097, 0.1004, 0.103, 0.105, 0.1066, 0.1092, 0.1129],
    [7, 0.0705, 0.0761, 0.0806, 0.0844, 0.0877, 0.0904, 0.0949, 0.0982, 0.1008, 0.1028, 0.1044, 0.1071, 0.1109],
    [7.2, 0.0688, 0.0742, 0.0787, 0.0825, 0.0857, 0.0884, 0.0928, 0.0962, 0.0987, 0.1008, 0.1023, 0.1051, 0.109],
    [7.4, 0.0672, 0.0725, 0.0769, 0.0806, 0.0838, 0.0865, 0.0908, 0.0942, 0.0967, 0.0988, 0.1004, 0.1031, 0.1071],
    [7.6, 0.0656, 0.0709, 0.0752, 0.0789, 0.082, 0.0846, 0.0889, 0.0922, 0.0948, 0.0968, 0.0984, 0.1012, 0.1054],
    [7.8, 0.0642, 0.0693, 0.0736, 0.0771, 0.0802, 0.0828, 0.0871, 0.0904, 0.0929, 0.095, 0.0966, 0.0994, 0.1036],
    [8, 0.0627, 0.0678, 0.072, 0.0755, 0.0785, 0.0811, 0.0853, 0.0886, 0.0912, 0.0932, 0.0948, 0.0976, 0.102],
    [8.2, 0.0614, 0.0663, 0.0705, 0.0739, 0.0769, 0.0795, 0.0837, 0.0869, 0.0894, 0.0914, 0.0931, 0.0959, 0.1004],
    [8.4, 0.0601, 0.0649, 0.069, 0.0724, 0.0754, 0.0779, 0.082, 0.0852, 0.0878, 0.0893, 0.0914, 0.0943, 0.0938],
    [8.6, 0.0588, 0.0636, 0.0676, 0.071, 0.0739, 0.0764, 0.0805, 0.0836, 0.0862, 0.0882, 0.0898, 0.0927, 0.0973],
    [8.8, 0.0576, 0.0623, 0.0663, 0.0696, 0.0724, 0.0749, 0.079, 0.0821, 0.0846, 0.0866, 0.0882, 0.0912, 0.0959],
    [9.2, 0.0554, 0.0559, 0.0637, 0.067, 0.0697, 0.0721, 0.0761, 0.0792, 0.0817, 0.0837, 0.0853, 0.0882, 0.0931],
    [9.6, 0.0533, 0.0577, 0.0614, 0.0645, 0.0672, 0.0696, 0.0734, 0.0765, 0.0789, 0.0809, 0.0825, 0.0855, 0.0905],
    [10, 0.0514, 0.0556, 0.0592, 0.0622, 0.0649, 0.0672, 0.071, 0.0739, 0.0763, 0.0783, 0.0799, 0.0829, 0.088],
    [10.4, 0.0496, 0.0537, 0.0572, 0.0601, 0.0627, 0.0649, 0.0686, 0.0716, 0.0739, 0.0759, 0.0775, 0.0804, 0.0857],
    [10.8, 0.0479, 0.0519, 0.0553, 0.0581, 0.0606, 0.0628, 0.0664, 0.0693, 0.0717, 0.0736, 0.0751, 0.0781, 0.0834],
    [11.2, 0.0463, 0.0502, 0.0552, 0.0563, 0.0587, 0.0609, 0.0644, 0.0672, 0.0695, 0.0714, 0.073, 0.0759, 0.0813],
    [11.6, 0.0448, 0.0486, 0.0518, 0.0545, 0.0569, 0.059, 0.0625, 0.0652, 0.0675, 0.0694, 0.0709, 0.0738, 0.0793],
    [12, 0.0435, 0.0471, 0.0502, 0.0529, 0.0552, 0.0573, 0.0606, 0.0634, 0.0656, 0.0674, 0.069, 0.0719, 0.0774],
    [12.8, 0.0409, 0.0444, 0.0474, 0.0499, 0.0521, 0.0541, 0.0573, 0.0599, 0.0621, 0.0639, 0.0654, 0.0682, 0.0739],
    [13.6, 0.0387, 0.042, 0.0448, 0.0472, 0.0493, 0.0512, 0.0543, 0.0568, 0.0589, 0.0607, 0.0621, 0.0649, 0.0707],
    [14.4, 0.0367, 0.0398, 0.0425, 0.0448, 0.0468, 0.0486, 0.0516, 0.054, 0.0561, 0.0577, 0.0592, 0.0619, 0.0677],
    [15.2, 0.0349, 0.0379, 0.0404, 0.0426, 0.0446, 0.0463, 0.0492, 0.0515, 0.0535, 0.0551, 0.0565, 0.0592, 0.065],
    [16, 0.0332, 0.0361, 0.0385, 0.0407, 0.0425, 0.0442, 0.0469, 0.0492, 0.0511, 0.0527, 0.054, 0.0567, 0.0625],
    [18, 0.0297, 0.0323, 0.0345, 0.0364, 0.0381, 0.0396, 0.0422, 0.0442, 0.046, 0.0475, 0.0487, 0.0512, 0.057],
    [20, 0.0269, 0.0292, 0.0312, 0.033, 0.0345, 0.0359, 0.0383, 0.0402, 0.0418, 0.0432, 0.0444, 0.0468, 0.0524]
]


# --- 帮助与计算函数 ---

def linear_interpolate(x1, y1, x2, y2, x):
    """一维线性插值"""
    if x1 == x2:
        return y1
    return y1 + (y2 - y1) * (x - x1) / (x2 - x1)


def bilinear_interpolate(table, z_b, l_b):
    """二维线性插值函数，用于 ALPHA_TABLE"""
    z_values = [float(row[0]) for row in table[1:] if row[0] is not None]
    l_values = [float(val) for val in table[0][1:] if val is not None]
    alpha_values = [[float(cell) if cell is not None else 0.0 for cell in row[1:]] for row in table[1:]]

    i = 0
    while i < len(z_values) - 1 and z_values[i] < z_b:
        i += 1
    i = max(0, min(i, len(z_values) - 2))

    j = 0
    while j < len(l_values) - 1 and l_values[j] < l_b:
        j += 1
    j = max(0, min(j, len(l_values) - 2))

    alpha11, alpha12 = alpha_values[i][j], alpha_values[i][j + 1]
    alpha21, alpha22 = alpha_values[i + 1][j], alpha_values[i + 1][j + 1]

    alpha1 = linear_interpolate(l_values[j], alpha11, l_values[j + 1], alpha12, l_b)
    alpha2 = linear_interpolate(l_values[j], alpha21, l_values[j + 1], alpha22, l_b)

    return linear_interpolate(z_values[i], alpha1, z_values[i + 1], alpha2, z_b)


def is_fill(layer_id):
    """判断地层是否为填土"""
    return layer_id and str(layer_id).startswith('1-')


# --- 数据解析函数 ---

def parse_buildings(sheet):
    """解析建筑物数据（1.1工作表）"""
    # 列索引（从0开始）
    NAME_COL = 2  # C列：建筑物名称
    FLOORS_COL = 4  # E列：层数（地上）
    HEIGHT_COL = 5  # F列：高度
    EMBED_ELEV_COL = 11  # L列：基础埋深标高
    LOAD_COL = 12  # M列：荷载
    WIDTH_COL = 13  # N列：宽度
    LENGTH_COL = 14  # O列：长度

    logging.info("开始解析建筑物数据（1.1工作表）")

    def parse_optional_float(val):
        if val is None or val == '' or str(val).strip() in {'/', '-', '—', '–', 'N/A'}:
            return None
        try:
            return float(val)
        except (ValueError, TypeError):
            return None

    buildings = {}
    buildings_list = []

    for row in sheet.iter_rows(min_row=6, values_only=True):
        if not row or len(row) <= NAME_COL or not row[NAME_COL]:
            continue

        name = str(row[NAME_COL]).strip()

        embed_elev_val = parse_optional_float(row[EMBED_ELEV_COL] if len(row) > EMBED_ELEV_COL else None)
        if embed_elev_val is None:
            logging.warning(f"建筑物 {name} 基础埋深标高无效，跳过")
            continue

        # === 修改点：解析层数（处理范围值和空值） ===
        floors_raw = row[FLOORS_COL] if len(row) > FLOORS_COL else None
        floors = 1  # 默认值为 1

        if floors_raw is not None:
            raw_str = str(floors_raw).strip()
            if raw_str in {'', '/', '-', '—', 'None'}:
                floors = 1
            else:
                try:
                    # 处理范围值 (例如 "10~12" 或 "10-12")
                    if '~' in raw_str:
                        parts = raw_str.split('~')
                        floors = max([float(p) for p in parts if p.strip()])
                    elif '-' in raw_str:  # 注意：只有当不是负数时才分割，简单起见假设层数都是正数范围
                        parts = raw_str.split('-')
                        floors = max([float(p) for p in parts if p.strip()])
                    else:
                        floors = float(raw_str)

                    floors = int(floors)  # 转为整数
                    if floors <= 0: floors = 1
                except (ValueError, TypeError):
                    floors = 1
        # ==========================================

        height = parse_optional_float(row[HEIGHT_COL] if len(row) > HEIGHT_COL else None)

        if height is None or height <= 0:
            if floors is not None and floors > 0:
                height = floors * 3.5
                logging.info(f"建筑物 {name} 高度缺失，使用层数 {floors} × 3.5m = {height}m 自动补全")
            else:
                logging.info(f"建筑物 {name} 高度和层数均无效，高度保持 None")

        width = parse_optional_float(row[WIDTH_COL] if len(row) > WIDTH_COL else None)
        length = parse_optional_float(row[LENGTH_COL] if len(row) > LENGTH_COL else None)
        load = parse_optional_float(row[LOAD_COL] if len(row) > LOAD_COL else None)

        buildings[name] = {
            'embed_elev': embed_elev_val,
            'width': width,
            'length': length,
            'height': height,
            'floors': floors,
            'load': load
        }
        buildings_list.append(name)

    return buildings, buildings_list


def parse_holes(sheet, buildings_list):
    """解析钻孔数据"""
    logging.info("开始解析钻孔数据（1.5单孔工作表）")
    holes = {}
    unmatched_buildings = set()
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row and row[0]:
            hole_id = str(row[0]).strip()

            def parse_optional_float(val):
                try:
                    return float(val) if val is not None else None
                except (ValueError, TypeError):
                    return None

            builds_str = row[12] if len(row) > 12 else ''
            assoc_builds = [b.strip() for b in re.split(r'[,、]', str(builds_str or '')) if b.strip()]

            for b in assoc_builds:
                if b and b not in buildings_list: unmatched_buildings.add(b)

            holes[hole_id] = {
                'elev': parse_optional_float(row[1]), 'max_depth': parse_optional_float(row[2]),
                'x': parse_optional_float(row[7]), 'y': parse_optional_float(row[8]),
                'builds': assoc_builds
            }
    if unmatched_buildings:
        logging.warning(f"以下建筑物在 1.5单孔 中未匹配到 1.1 的建筑物: {unmatched_buildings}")
    return holes


# ====================== 新增：全局样式设置函数 ======================
def force_set_all_fonts(doc):
    """强制设置文档中所有段落的字体（12pt） - 解决Word主题干扰的终极方案"""
    print("开始强制设置所有段落的字体...")

    # 1. 强制设置所有段落的字体
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # 只处理非空段落
            for run in paragraph.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                run.font.size = Pt(12)
                run.font.bold = False

    # 2. 强制设置所有表格的字体
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():  # 只处理非空段落
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                            # 表格使用五号字体(10.5pt)
                            run.font.size = Pt(10.5)

    print("强制设置字体完成！")


def set_global_styles(doc):
    """设置整个文档的全局样式（标题、正文、段落、字体等） - 终极修复版"""
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        # 1. 页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)  # 上 2.54cm
            section.bottom_margin = Inches(1.0)  # 下 2.54cm
            section.left_margin = Inches(1.0)  # 左 2.54cm
            section.right_margin = Inches(0.8)  # 右 2.0cm

        styles = doc.styles

        # 正文 Normal → 小四 宋体（12pt）- 终极100%生效版
        normal = styles['Normal']
        normal.font.name = '宋体'
        normal._element.rPr.rFonts.set(qn('w:ascii'), '宋体')  # 修复英文11pt
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 中文
        normal._element.rPr.rFonts.set(qn('w:cs'), '宋体')  # 可选：复杂脚本

        # 关闭主题字体干扰（关键修复）
        try:
            normal._element.rPr.append(parse_xml(
                f'<w:rPr {nsdecls("w")}><w:noProof/></w:rPr>'
            ))
        except:
            pass  # 忽略XML解析错误

        # 关闭主题字体属性
        try:
            themeFonts = normal.element.xpath('.//w:themeFont')
            if themeFonts:
                themeFonts[0].set('val', '')
        except:
            pass  # 忽略主题字体设置错误

        normal.font.size = Pt(12)  # 真正的12pt
        normal.font.bold = False
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.first_line_indent = Inches(0.74)

        # 一级标题 Heading 1 → 黑体 小三 加粗 居中
        h1 = styles['Heading 1']
        h1.font.name = '黑体'
        h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)

        # 二级标题 Heading 2 → 黑体 四号 加粗
        h2 = styles['Heading 2']
        h2.font.name = '黑体'
        h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)

        # 三级标题 Heading 3 → 黑体 小四 加粗
        h3 = styles['Heading 3']
        h3.font.name = '黑体'
        h3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h3.font.size = Pt(12)
        h3.font.bold = True

        # 表格样式（五号宋体）
        if 'Table Grid' in styles:
            table_style = styles['Table Grid']
            table_style.font.name = '宋体'
            table_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            table_style.font.size = Pt(10.5)

        # 可选：强制所有段落首行缩进（更保险）
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') or 'Table' in paragraph.style.name:
                continue
            paragraph.paragraph_format.first_line_indent = Inches(0.74)

        print("全局样式设置完成！（终极修复版）")
    except Exception as e:
        print(f"设置全局样式时出错: {e}")
        # 即使出错也要继续，确保不会因为样式问题中断整个流程
# ==================================================================

# ====================== 新增：全局表格字体函数（放在这里！）======================
def set_table_font(table):
    """为整个表格的所有单元格设置五号宋体（10.5pt）"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def parse_layer_info(sheet):
    """解析地层信息（1.6地层信息工作表）"""
    logging.info("开始解析地层信息（1.6地层信息工作表）")
    layer_info = {}
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row and row[0]:
            layer_id = str(row[0]).strip()

            # === 修改点：遇到 END 则跳过 ===
            if layer_id.upper() == 'END':
                continue
            # ==========================

            try:
                def safe_float(val, default=0.0):
                    if val is None or val == '' or str(val).strip() in {'-', '/', '—', '－', '无', '暂无'}:
                        return default
                    try:
                        return float(str(val).strip().replace(',', ''))
                    except:
                        return default

                name = str(row[7]).strip().replace(' ', '').replace('\u3000', '') if len(row) > 7 and row[7] else ''
                state = str(row[8]).strip().replace(' ', '').replace('\u3000', '') if len(row) > 8 and row[8] else '/'
                layer_info[layer_id] = {
                    'name': name,
                    'state': state,
                    'bearing_capacity': safe_float(row[2]),
                    'compression_modulus': safe_float(row[3]),
                    'density': safe_float(row[9])
                }
                logging.info(
                    f"Parsed layer {layer_id}: name={repr(name)}, fak={layer_info[layer_id]['bearing_capacity']}")
            except Exception as e:
                logging.warning(f"地层 {layer_id} 数据转换失败: {e}")
    return layer_info


def parse_hole_strata(sheet, holes):
    """解析各孔地层数据"""
    logging.info("开始解析各孔地层数据（2.4各孔地层工作表）")
    hole_strata = defaultdict(list)
    current_hole = None
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row and str(row[0]).strip().upper() == 'END':
            break

        if row and row[0]:
            current_hole = str(row[0]).strip()

        if current_hole and row and row[1]:
            final_depth = None
            try:
                layer_id = str(row[1]).strip()
                raw_depth = row[2] if len(row) > 2 else None

                if raw_depth is not None and str(raw_depth).strip() != '':
                    parsed_depth = float(raw_depth)
                    if parsed_depth > 0:
                        final_depth = parsed_depth
                    else:
                        logging.warning(f"钻孔 {current_hole} 层 {layer_id} 深度值 {parsed_depth} 无效。")

                if final_depth is None:
                    max_depth = holes.get(current_hole, {}).get('max_depth')
                    if max_depth is not None and max_depth > 0:
                        final_depth = max_depth
                        logging.info(f"钻孔 {current_hole} 层 {layer_id} 深度为空或无效，使用孔深 {final_depth}m。")
                    else:
                        logging.warning(f"钻孔 {current_hole} 层 {layer_id} 无有效层底深度或孔深，跳过。")

                if final_depth is not None:
                    hole_strata[current_hole].append((layer_id, final_depth))

            except (ValueError, TypeError) as e:
                logging.warning(f"钻孔 {current_hole} 地层数据转换失败: {e}，跳过该行。")
    return dict(hole_strata)


# --- 第1节：地基均匀性评价相关函数 ---

def find_layer_at_depth(hole_strata, holes, hole_id, depth):
    layers = hole_strata.get(hole_id, [])
    if not layers: return None
    prev_bottom = 0.0
    for layer_id, bottom in layers:
        if prev_bottom <= depth <= bottom:
            return layer_id
        prev_bottom = bottom
    return layers[-1][0] if layers else None


def compute_desc(building, buildings, holes, layer_info, hole_strata):
    logging.info(f"计算建筑物 {building} 的地层描述")
    build_holes_list = [h for h, info in holes.items() if building in info['builds']]
    logging.info(f"建筑物 {building} 关联钻孔列表: {build_holes_list}")
    if not build_holes_list:
        logging.warning(f"建筑物 {building} 无关联钻孔 - 请检查1.5单孔中的名称是否与1.1的C列完全匹配（无多余空格、逗号等）")
        return "无勘探点数据"

    hole_elevs = [holes[h]['elev'] for h in build_holes_list if holes[h]['elev'] is not None]
    if not hole_elevs:
        logging.warning(f"建筑物 {building} 无有效钻孔标高数据")
        return "无有效标高数据"
    min_elev = min(hole_elevs)
    base_elev = buildings[building]['embed_elev']
    max_above = base_elev - min_elev

    hole_to_layer = {}
    fill_thicks = defaultdict(list)
    first_layers = []
    above_holes = 0

    for hole_id in build_holes_list:
        hole_elev = holes[hole_id]['elev']
        if hole_elev is None:
            logging.warning(f"钻孔 {hole_id} 无标高数据，跳过")
            continue
        embed_depth = hole_elev - base_elev
        strata = hole_strata[hole_id]
        if not strata:
            logging.warning(f"钻孔 {hole_id} 无地层数据，跳过")
            continue
        first_layer = strata[0][0]
        first_layers.append(first_layer)

        if embed_depth < 0:
            above_holes += 1
            logging.info(f"钻孔 {hole_id} 基底高于地面，嵌入深度 {embed_depth}")
            continue

        layer = find_layer_at_depth(hole_strata, holes, hole_id, embed_depth)
        if layer:
            hole_to_layer[hole_id] = layer
            if is_fill(layer):
                layer_bottom = 0.0
                for l, b in strata:
                    if not is_fill(l):
                        break
                    layer_bottom = b if b is not None else holes[hole_id]['max_depth'] or 0.0
                if embed_depth is not None and layer_bottom is not None:
                    fill_thick = layer_bottom - embed_depth
                    fill_thicks[layer].append(fill_thick)
                    if fill_thick <= 0:
                        logging.warning(f"钻孔 {hole_id} 计算填土厚度为非正值（{fill_thick}），仍记录")
                    logging.info(f"钻孔 {hole_id} 基底为填土层 {layer}，嵌入深度={embed_depth}, 层底深度={layer_bottom}, 填土厚度={fill_thick}")
                else:
                    logging.warning(f"钻孔 {hole_id} 无法计算填土厚度：嵌入深度={embed_depth}, 层底深度={layer_bottom}")

    desc = ""
    if above_holes > 0 or max_above >= 0:
        max_above = round(max(0, max_above), 2)
        first_counts = Counter(first_layers)
        if first_counts:
            main_first, _ = first_counts.most_common(1)[0]
            main_first_name = layer_info.get(main_first, {}).get('name', '')
            desc += f"基底高于现状地面最大约{max_above}m，现状地面下主要分布{main_first_name}{main_first}"
            logging.info(f"建筑物 {building} 基底高于地面，表面地层 {main_first_name}{main_first}")
        else:
            desc += f"基底高于现状地面最大约{max_above}m，现状地面下无地层信息"
            logging.warning(f"建筑物 {building} 基底高于地面，无表面地层信息")
        if hole_to_layer:
            desc += "；"
        else:
            if first_layers and is_fill(first_layers[0]) and fill_thicks.get(first_layers[0]):
                positive_thicks = [t for t in fill_thicks[first_layers[0]] if t > 0]
                if positive_thicks:
                    fill_thick = round(max(positive_thicks), 2)
                    desc += f"；基底分布地层为最大{fill_thick}m厚{main_first_name}{first_layers[0]}"
                    logging.info(f"建筑物 {building} 基底为填土层，最大厚度 {fill_thick}")
                else:
                    desc += f"；基底分布地层为{main_first_name}{first_layers[0]}，无有效填土厚度"
                    logging.warning(f"建筑物 {building} 基底为填土层，无正值填土厚度")
            else:
                desc += f"；基底分布地层为{main_first_name}{first_layers[0]}"
            desc += "。"
            return desc

    if hole_to_layer:
        layer_counts = Counter(hole_to_layer.values())
        if layer_counts:
            main_layer, main_count = layer_counts.most_common(1)[0]
            main_name = layer_info.get(main_layer, {}).get('name', '')
            main_desc = f"{main_name}{main_layer}"
            if is_fill(main_layer) and fill_thicks[main_layer]:
                positive_thicks = [t for t in fill_thicks[main_layer] if t > 0]
                if positive_thicks:
                    mt = round(max(positive_thicks), 2)
                    main_desc = f"最大{mt}m厚{main_desc}"
            if len(layer_counts) == 1:
                desc += f"基底分布地层为{main_desc}"
                logging.info(f"建筑物 {building} 单一地层: {main_desc}")
            else:
                desc += f"基底主要分布地层为{main_desc}"
                logging.info(f"建筑物 {building} 主要地层: {main_desc}")

            others = layer_counts.most_common()[1:]
            if others:
                total_holes = len(hole_to_layer)
                prop_others = (total_holes - main_count) / total_holes
                word = "个别" if prop_others < 0.3 else "部分"
                other_descs = []
                for l, _ in others:
                    n = layer_info.get(l, {}).get('name', '')
                    od = f"{n}{l}"
                    if is_fill(l) and fill_thicks[l]:
                        positive_thicks = [t for t in fill_thicks[l] if t > 0]
                        if positive_thicks:
                            mt = round(max(positive_thicks), 2)
                            od = f"最大{mt}m厚{od}"
                        else:
                            od = f"{od}（无有效填土厚度）"
                    other_descs.append(od)
                desc += f"，{word}地段分布{'、'.join(other_descs)}"
                logging.info(f"建筑物 {building} 次要地层: {other_descs}")

        desc += "。"

    if not desc:
        desc = "无有效数据"
        logging.warning(f"建筑物 {building} 无有效地层描述")
    return desc


def is_high_rise(building, buildings):
    info = buildings.get(building, {})
    floors, height = info.get('floors'), info.get('height')
    is_res = any(k in building.lower() for k in ['住宅', '公寓', '宿舍'])
    if floors is not None and floors >= 7: return True
    if height is not None and ((is_res and height >= 27) or (not is_res and height >= 24)): return True
    return False


def calculate_effective_fill_thick(hole_id, hole_strata, holes, base_elev):
    hole = holes[hole_id]
    if hole['elev'] is None: return 0.0
    embed_depth = hole['elev'] - base_elev

    strata = hole_strata.get(hole_id, [])

    if embed_depth < 0:
        ground_fill_thick = 0.0
        for l_id, bottom in strata:
            if not is_fill(l_id): break
            ground_fill_thick = bottom
        return -embed_depth + ground_fill_thick
    else:
        fill_bottom, found_base_in_fill = embed_depth, False
        for l_id, bottom in strata:
            if bottom < embed_depth: continue
            if not found_base_in_fill:
                if not is_fill(find_layer_at_depth(hole_strata, holes, hole_id, embed_depth)): return 0.0
                found_base_in_fill = True

            if is_fill(l_id):
                fill_bottom = bottom
            else:
                break
        return fill_bottom - embed_depth


def needs_equivalent_modulus(building, holes, hole_strata, buildings):
    """
    判断是否需要进行当量模量计算
    条件：
    1. 是高层建筑
    2. 基底无填土 OR (基底有填土 且 填土满足换填标准)
    """
    # 1. 首先必须是高层
    if not is_high_rise(building, buildings):
        return False, False  # (是否计算, 是否涉及换填)

    build_holes = [h for h, i in holes.items() if building in i.get('builds', [])]
    if not build_holes:
        return False, False

    base_elev = buildings[building]['embed_elev']

    # 收集该楼栋下所有孔的填土厚度
    fill_thicknesses = []
    has_fill = False

    for hole_id in build_holes:
        # 计算该孔基底以下的填土厚度
        th = calculate_effective_fill_thick(hole_id, hole_strata, holes, base_elev)
        if th > 0:
            has_fill = True
        fill_thicknesses.append(th)

    if not has_fill:
        # 情况A：完全无填土 -> 需要计算，不涉及换填
        return True, False

    # 情况B/C：有填土，判断是否满足换填标准
    # 标准：
    # 1. 所有孔厚度 < 1m  -> 满足
    # OR
    # 2. 1m<=厚度<=2m 的孔数 <= 2 AND 2m<厚度<=3m 的孔数 <= 1 AND >3m 的孔数 == 0

    # 统计各区间孔数
    c_gt_3 = sum(1 for t in fill_thicknesses if t > 3.0)

    if c_gt_3 > 0:
        return False, False  # 有厚度>3m的孔，不可换填 -> 不计算

    # 如果最大厚度 < 1m -> 直接满足
    if max(fill_thicknesses) < 1.0:
        return True, True  # 需要计算，且涉及换填

    # 复杂条件判断
    c_1_to_2 = sum(1 for t in fill_thicknesses if 1.0 <= t <= 2.0)
    c_2_to_3 = sum(1 for t in fill_thicknesses if 2.0 < t <= 3.0)

    if c_1_to_2 <= 2 and c_2_to_3 <= 1:
        return True, True  # 满足换填标准 -> 需要计算，且涉及换填

    return False, False


def calculate_effective_embed_depth(hole_id, hole_strata, holes, base_elev):
    hole = holes[hole_id]
    if hole['elev'] is None: return 0.0
    embed_depth = hole['elev'] - base_elev

    base_layer = find_layer_at_depth(hole_strata, holes, hole_id, max(0, embed_depth))
    if not is_fill(base_layer): return max(0, embed_depth)

    current_depth = max(0, embed_depth)
    for l_id, bottom in hole_strata.get(hole_id, []):
        if bottom < current_depth: continue
        if is_fill(l_id):
            current_depth = bottom
        else:
            break
    return current_depth


def compute_equivalent_modulus(building, holes, hole_strata, buildings, layer_info, is_replace_mode=False):
    """
    计算当量模量
    is_replace_mode: 是否为换填模式。如果是，计算深度从填土底面开始。
    """
    build_holes_list = [h for h, info in holes.items() if building in info.get('builds', [])]
    if not build_holes_list: return None

    b_info = buildings[building]
    width, length = b_info.get('width'), b_info.get('length')
    if not all((width, length)): return None

    b, l_b = width / 2, length / width
    total_numerator, total_denominator, es_values = 0.0, 0.0, []

    base_elev = b_info['embed_elev']

    for hole_id in build_holes_list:
        if holes[hole_id].get('elev') is None: continue

        # 1. 确定计算起始深度 (相对于自然地面)
        # 默认起始深度 = 基础埋深
        hole_elev = holes[hole_id]['elev']
        embed_depth = hole_elev - base_elev

        # 如果是换填模式，且该孔有填土，起始深度下移到填土底
        # 注意：calculate_effective_embed_depth 原本就是找填土底，正好复用！
        # 如果不是换填模式（即无填土），它也会返回基础埋深
        # 所以这里直接用 calculate_effective_embed_depth 即可满足两种情况
        # 因为根据前一步的筛选，能进来的有填土楼栋都是“可换填”的。

        if is_replace_mode:
            # 换填模式：起始计算深度 = 填土层底面 (忽略填土变形)
            effective_start_depth = calculate_effective_embed_depth(hole_id, hole_strata, holes, base_elev)
        else:
            # 无填土模式：起始计算深度 = 基础埋深
            # 防止浮点误差，取max(0)
            effective_start_depth = max(0.0, embed_depth)

        hole_depth = holes[hole_id].get('max_depth')
        if hole_depth is None: continue

        # 2. 收集起始深度以下的地层
        layers_below, prev_depth = [], 0
        for layer_id, bottom in hole_strata.get(hole_id, []):
            effective_bottom = bottom if bottom is not None else hole_depth

            if effective_bottom is None: continue

            # 只计算起始深度以下的层
            if effective_bottom > effective_start_depth:
                # 该层的有效计算段：从 max(层顶, 起始深度) 到 层底
                start = max(effective_start_depth, prev_depth)

                # 排除填土层（如果是换填模式，填土层应该都在 start 之上，但为了保险再次检查）
                if is_replace_mode and is_fill(layer_id):
                    pass  # 换填模式下，不计算填土层的变形
                else:
                    if start < effective_bottom:
                        layers_below.append((layer_id, start, effective_bottom))

            prev_depth = effective_bottom

        if not layers_below: continue

        hole_numerator, hole_denominator, alpha_prev, zi_prev = 0.0, 0.0, 0.0, 0.0

        # 3. 开始分层求和
        # zi 的零点是 effective_start_depth (即基底或换填底面)
        for layer_id, start, bottom in layers_below:
            zi = bottom - effective_start_depth
            z_b = zi / b if b != 0 else float('inf')
            alpha = bilinear_interpolate(ALPHA_TABLE, z_b, l_b)

            delta_sigma = alpha * zi - alpha_prev * zi_prev
            comp_mod = layer_info.get(layer_id, {}).get('compression_modulus')
            # 容错：如果没有模量，默认为一个较大值或者跳过（防止除零）
            if comp_mod is None or comp_mod == 0:
                # logging.warning(f"{hole_id} {layer_id} 模量为0，跳过该层计算")
                continue

            if delta_sigma != 0:
                hole_numerator += delta_sigma
                hole_denominator += delta_sigma / comp_mod

            alpha_prev, zi_prev = alpha, zi

        if hole_denominator != 0:
            es_hole = hole_numerator / hole_denominator
            es_values.append(es_hole)
            total_numerator += hole_numerator
            total_denominator += hole_denominator

    if not es_values or total_denominator == 0: return None

    es_building = total_numerator / total_denominator
    es_values.append(es_building)
    es_max, es_min = max(es_values), min(es_values)
    es_avg = sum(es_values) / len(es_values)
    es_ratio = es_max / es_min if es_min != 0 else float('inf')

    if es_avg <= 4:
        k = 1.3
    elif es_avg <= 7.5:
        k = linear_interpolate(4, 1.3, 7.5, 1.5, es_avg)
    elif es_avg <= 15:
        k = linear_interpolate(7.5, 1.5, 15, 1.8, es_avg)
    elif es_avg < 20:
        k = linear_interpolate(15, 1.8, 20, 2.5, es_avg)
    else:
        k = 2.5

    return {'es_max': round(es_max, 2), 'es_min': round(es_min, 2), 'es_avg': round(es_avg, 2),
            'es_ratio': round(es_ratio, 2), 'k': round(k, 2), 'uniformity': '均匀' if es_ratio <= k else '不均匀'}

# --- 第2节：天然地基评价相关函数 ---

def get_base_layers(building, buildings, holes, hole_strata, layer_info):
    """修改后的函数，返回更多信息包括高出地面情况"""
    base_layers = set()
    above_ground_heights = []  # 收集高出高度以计算平均
    above_ground_info = []  # 存储高出地面的地层信息和承载力满足情况
    build_holes = [h for h, info in holes.items() if building in info['builds']]
    all_above_ground = True
    load_kpa = buildings[building].get('load', None)

    for hole_id in build_holes:
        if hole_id not in hole_strata:
            logging.warning(f"钻孔 {hole_id} 无地层数据，跳过")
            continue
        hole_elev = holes[hole_id]['elev']
        if hole_elev is None:
            logging.warning(f"钻孔 {hole_id} 无标高数据，跳过")
            continue
        base_elev = buildings[building]['embed_elev']
        embed_depth = hole_elev - base_elev
        strata = hole_strata[hole_id]

        if embed_depth < 0:  # 基底高于地面
            if strata:
                layer_id = strata[0][0]  # 地表层
                layer_name = layer_info.get(layer_id, {}).get('name', '')
                high = -embed_depth
                above_ground_heights.append(high)
                base_layers.add(f"{layer_name}{layer_id}（基底标高高于钻孔标高 {high:.2f}m）")

                # 计算承载力满足情况
                if load_kpa is not None:
                    fak = layer_info.get(layer_id, {}).get('bearing_capacity')
                    if fak is not None:
                        ind_satisfy = '满足' if fak > load_kpa else '不满足'
                        raft_satisfy = '满足' if fak > load_kpa else '不满足'
                        above_ground_info.append((f"{layer_name}{layer_id}", ind_satisfy, raft_satisfy))
            continue
        else:
            all_above_ground = False

        # 正常情况：基底在地下
        found = False
        for layer_id, bottom in strata:
            effective_bottom = bottom if bottom is not None else holes[hole_id].get('max_depth', None)
            if effective_bottom is None or effective_bottom <= 0:
                continue
            if embed_depth <= effective_bottom:
                layer_name = layer_info.get(layer_id, {}).get('name', '')
                base_layers.add(f"{layer_name}{layer_id}")
                found = True
                break
        if not found:
            logging.warning(f"钻孔 {hole_id} 未找到基底层")

    avg_above_height = sum(above_ground_heights) / len(above_ground_heights) if above_ground_heights else 0
    return sorted(list(base_layers)), all_above_ground, avg_above_height, above_ground_info


def get_eta_params(layer_name, silt_params, clay_params):
    layer_name_lower = layer_name.lower().strip()
    if '淤泥' in layer_name_lower or '淤泥质土' in layer_name_lower:
        return 0.0, 1.0, False
    elif any(word in layer_name_lower for word in ['粉砂', '细砂']) or (
            '全风化' in layer_name_lower and '泥岩' not in layer_name_lower) or (
            '强风化' in layer_name_lower and '泥岩' not in layer_name_lower):
        return 2.0, 3.0, False
    elif ('全风化' in layer_name_lower or '强风化' in layer_name_lower) and '泥岩' in layer_name_lower:
        return 0.3, 1.6, False
    elif '中风化' in layer_name_lower or '中等风化' in layer_name_lower or '微风化' in layer_name_lower:
        return 0.0, 0.0, False
    elif '粉土' in layer_name_lower:
        eta_b, eta_d = silt_params.get('粉土', (0.5, 2.0))
        return eta_b, eta_d, False
    elif '粉质黏土' in layer_name_lower or '粉质粘土' in layer_name_lower:
        eta_b, eta_d = clay_params.get('粉质黏土', (0.3, 1.6))
        return eta_b, eta_d, False
    elif '黏土' in layer_name_lower or '粘土' in layer_name_lower:
        return 0.0, 1.0, True
    elif any(word in layer_name_lower for word in
             ['中砂', '粗砂', '砾砂', '漂石', '块石', '卵石', '砂卵石', '碎石', '圆砾', '角砾']):
        return 3.0, 4.4, False
    else:
        logging.warning(f"未知地层类型 {layer_name}，使用ηb=0, ηd=0")
        return 0.0, 0.0, False


def calculate_fa(fak, eta_b, eta_d, is_clay, b, d, layer_name, gamma, gamma_m):
    if fak is None or fak == 0: return "/"
    if '中风化' in layer_name or '微风化' in layer_name: return fak
    d_term = (d - 1) if is_clay else (d - 0.5)
    return round(fak + eta_b * gamma * (b - 3) + eta_d * gamma_m * d_term)


def get_weak_underlayers(building, buildings, holes, hole_strata, layer_info):
    """获取软弱下卧层信息，并预计算必要的参数"""
    load_kpa = buildings[building].get('load')
    if load_kpa is None: return []

    weak_layers, base_elev = [], buildings[building]['embed_elev']
    for hole_id in [h for h, i in holes.items() if building in i.get('builds', [])]:
        hole = holes[hole_id]
        strata = hole_strata.get(hole_id)
        if hole.get('elev') is None or not strata: continue

        embed_depth = hole['elev'] - base_elev
        if embed_depth < 0: continue

        base_layer_id = find_layer_at_depth(hole_strata, holes, hole_id, embed_depth)
        if not base_layer_id: continue

        found_base, prev_depth = False, 0
        for l_id, bottom in strata:
            if not found_base and embed_depth <= bottom:
                found_base, prev_depth = True, bottom
                continue
            if found_base:
                fak = layer_info.get(l_id, {}).get('bearing_capacity')
                if fak is not None and fak < load_kpa:
                    es1 = layer_info.get(base_layer_id, {}).get('compression_modulus')
                    es2 = layer_info.get(l_id, {}).get('compression_modulus')
                    D = bottom - prev_depth  # 下卧层厚度
                    weak_layers.append({
                        'layer_str': f"{layer_info.get(l_id, {}).get('name', '')}{l_id}",
                        'Z': prev_depth - embed_depth,  # 基底到下卧层顶面的距离
                        'D': D,
                        'weak_layer_id': l_id,
                        'es1_es2': (es1 / es2) if es1 and es2 else 10,
                        'fak': fak,  # 存储原始承载力
                        'layer_name': layer_info.get(l_id, {}).get('name', '')  # 存储地层名称用于计算ηd
                    })
                prev_depth = bottom
    return weak_layers

def calculate_theta(z_b, es1_es2):
    if z_b < 0.25: return 0
    table = {3: (6, 23), 5: (10, 25), 10: (20, 30)}
    keys = sorted(table.keys())

    t0, t1 = table[keys[0]]
    if es1_es2 >= keys[-1]:
        t0, t1 = table[keys[-1]]
    else:
        for i in range(len(keys) - 1):
            if keys[i] <= es1_es2 <= keys[i + 1]:
                es1, es2 = keys[i], keys[i + 1]
                t0 = linear_interpolate(es1, table[es1][0], es2, table[es2][0], es1_es2)
                t1 = linear_interpolate(es1, table[es1][1], es2, table[es2][1], es1_es2)
                break

    if z_b >= 0.5: return t1
    return linear_interpolate(0.25, t0, 0.5, t1, z_b)

def get_under_layers(hole_id, embed_depth, hole_strata, layer_info):
    strata = hole_strata.get(hole_id, [])
    under_layers = []
    found_base = False
    base_layer_id = None
    prev_bottom = 0.0  # from ground
    for layer_id, bottom in strata:
        effective_bottom = bottom if bottom is not None else float('inf')
        layer_top = prev_bottom
        if not found_base:
            if embed_depth <= effective_bottom:
                found_base = True
                base_layer_id = layer_id
                esi = layer_info.get(layer_id, {}).get('compression_modulus', None)
                if esi:
                    z_i = effective_bottom - embed_depth
                    if z_i > 0:
                        under_layers.append({
                            'layer_id': layer_id,
                            'z_i': z_i,
                            'z_im1': 0,
                            'esi': esi
                        })
                prev_bottom = effective_bottom
                continue
        if found_base:
            esi = layer_info.get(layer_id, {}).get('compression_modulus', None)
            if esi:
                z_i = effective_bottom - embed_depth
                z_im1 = prev_bottom - embed_depth
                if z_i > z_im1:
                    under_layers.append({
                        'layer_id': layer_id,
                        'z_i': z_i,
                        'z_im1': z_im1,
                        'esi': esi
                    })
            prev_bottom = effective_bottom
    return under_layers, base_layer_id





def calculate_settlement(building, hole_id, buildings, holes, hole_strata, layer_info, alpha_data, z_b_values,
                         l_b_values):
    B = buildings[building].get('width', None)
    L = buildings[building].get('length', None)
    pk = buildings[building].get('load', None)
    if B is None or L is None or pk is None or B == 0 or L == 0:
        logging.warning(f"建筑物 {building} 缺少 B={B}, L={L} 或 pk={pk}，沉降计算返回0")
        return 0.0

    base_elev = buildings[building]['embed_elev']
    hole_elev = holes[hole_id]['elev']
    embed_depth = hole_elev - base_elev
    if embed_depth < 0:
        logging.warning(f"钻孔 {hole_id} 嵌入深度负，沉降计算返回0")
        return 0.0

    under_layers, base_layer_id = get_under_layers(hole_id, embed_depth, hole_strata, layer_info)
    if not under_layers:
        logging.info(f"钻孔 {hole_id} 无下卧层，沉降=0")
        return 0.0

    fak_base = layer_info.get(base_layer_id, {}).get('bearing_capacity', None)
    if fak_base is None:
        logging.warning(f"基底层 {base_layer_id} 无fak，沉降计算返回0")
        return 0.0

    l_over_b = L / B if B != 0 else 1.0  # 避免除零
    sum_numerator = 0.0
    sum_denominator = 0.0
    sum_settlement_terms = 0.0

    alpha_im1 = 0.0
    z_im1 = 0.0  # z0 = 0
    for layer in under_layers:
        z_i = layer['z_i']
        esi = layer['esi']
        # 计算 alpha_i
        z_b_i = z_i / (B / 2) if B != 0 else 0.0
        alpha_i = interpolate_alpha(alpha_data, z_b_values, l_b_values, z_b_i, l_over_b)

        term = alpha_i * z_i - alpha_im1 * z_im1
        sum_numerator += term
        sum_denominator += term / esi if esi != 0 else 0.0

        settlement_term = pk * 4 * term / esi if esi != 0 else 0.0
        sum_settlement_terms += settlement_term

        alpha_im1 = alpha_i
        z_im1 = z_i

    if sum_denominator == 0:
        logging.warning(f"sum_denominator=0，沉降计算返回0")
        return 0.0

    Es = sum_numerator / sum_denominator

    # 确定ψs
    es_values = [1.5, 2.5, 4.0, 7.0, 20.0]
    if pk >= fak_base:
        ps_values = [0.4, 1.4, 1.3, 1.0, 0.2]
    elif pk <= 0.75 * fak_base:
        ps_values = [0.4, 1.1, 1.0, 0.7, 0.2]
    else:
        # 线性插值 ps
        ps_low = [0.4, 1.1, 1.0, 0.7, 0.2]
        ps_high = [0.4, 1.4, 1.3, 1.0, 0.2]
        factor = (pk - 0.75 * fak_base) / (fak_base - 0.75 * fak_base)
        ps_values = [low + factor * (high - low) for low, high in zip(ps_low, ps_high)]

    if Es <= es_values[0]:
        ps = ps_values[0]
    elif Es >= es_values[-1]:
        ps = ps_values[-1]
    else:
        for i in range(len(es_values) - 1):
            if es_values[i] <= Es <= es_values[i + 1]:
                ps = linear_interpolate(es_values[i], ps_values[i], es_values[i + 1], ps_values[i + 1], Es)
                break

    settlement = ps * sum_settlement_terms
    logging.info(f"建筑物 {building} 钻孔 {hole_id} 沉降计算: Es={Es}, ψs={ps}, 沉降={settlement}")
    return settlement


def get_tilt_limit(height):
    if height is None:
        return 0.004
    if str(height).strip() in {'', '-', '/', '—', '－', '无', '暂无', 'None', 'null', 'NULL', 'N/A'}:
        return 0.004
    if height <= 24: return 0.004
    if height <= 60: return 0.003
    if height <= 100: return 0.0025
    return 0.002


# --- Word 文档生成函数 ---
def get_alpha_data():
    alpha_data = {
        0: {1: 0.25, 1.2: 0.25, 1.4: 0.25, 1.6: 0.25, 1.8: 0.25, 2: 0.25, 2.4: 0.25, 2.8: 0.25, 3.2: 0.25, 3.6: 0.25,
            4: 0.25, 5: 0.25, 10: 0.25},
        0.2: {1: 0.2496, 1.2: 0.2497, 1.4: 0.2497, 1.6: 0.2498, 1.8: 0.2498, 2: 0.2498, 2.4: 0.2498, 2.8: 0.2498,
              3.2: 0.2498, 3.6: 0.2498, 4: 0.2498, 5: 0.2498, 10: 0.2498},
        0.4: {1: 0.2474, 1.2: 0.2479, 1.4: 0.2481, 1.6: 0.2483, 1.8: 0.2483, 2: 0.2484, 2.4: 0.2485, 2.8: 0.2485,
              3.2: 0.2485, 3.6: 0.2485, 4: 0.2485, 5: 0.2485, 10: 0.2485},
        0.6: {1: 0.2423, 1.2: 0.2437, 1.4: 0.2444, 1.6: 0.2449, 1.8: 0.2451, 2: 0.2452, 2.4: 0.2454, 2.8: 0.2455,
              3.2: 0.2455, 3.6: 0.2455, 4: 0.2455, 5: 0.2455, 10: 0.2456},
        0.8: {1: 0.2346, 1.2: 0.2372, 1.4: 0.2387, 1.6: 0.2395, 1.8: 0.24, 2: 0.2403, 2.4: 0.2407, 2.8: 0.2408,
              3.2: 0.2409, 3.6: 0.2409, 4: 0.241, 5: 0.241, 10: 0.241},
        1: {1: 0.2252, 1.2: 0.2291, 1.4: 0.2313, 1.6: 0.2326, 1.8: 0.2335, 2: 0.234, 2.4: 0.2346, 2.8: 0.2349,
            3.2: 0.2351, 3.6: 0.2352, 4: 0.2352, 5: 0.2353, 10: 0.2353},
        1.2: {1: 0.2149, 1.2: 0.2199, 1.4: 0.2229, 1.6: 0.2248, 1.8: 0.226, 2: 0.2268, 2.4: 0.2278, 2.8: 0.2282,
              3.2: 0.2285, 3.6: 0.2286, 4: 0.2287, 5: 0.2288, 10: 0.2289},
        1.4: {1: 0.2043, 1.2: 0.2102, 1.4: 0.214, 1.6: 0.2146, 1.8: 0.218, 2: 0.2191, 2.4: 0.2204, 2.8: 0.2211,
              3.2: 0.2215, 3.6: 0.2217, 4: 0.2218, 5: 0.222, 10: 0.2221},
        1.6: {1: 0.1939, 1.2: 0.2006, 1.4: 0.2049, 1.6: 0.2079, 1.8: 0.2099, 2: 0.2113, 2.4: 0.213, 2.8: 0.2138,
              3.2: 0.2143, 3.6: 0.2146, 4: 0.2148, 5: 0.215, 10: 0.2152},
        1.8: {1: 0.184, 1.2: 0.1912, 1.4: 0.196, 1.6: 0.1994, 1.8: 0.2018, 2: 0.2034, 2.4: 0.2055, 2.8: 0.2066,
              3.2: 0.2073, 3.6: 0.2077, 4: 0.2079, 5: 0.2082, 10: 0.2084},
        2: {1: 0.1746, 1.2: 0.1822, 1.4: 0.1875, 1.6: 0.1912, 1.8: 0.198, 2: 0.1958, 2.4: 0.1982, 2.8: 0.1996,
            3.2: 0.2004, 3.6: 0.2009, 4: 0.2012, 5: 0.2015, 10: 0.2018},
        2.2: {1: 0.1659, 1.2: 0.1737, 1.4: 0.1793, 1.6: 0.1883, 1.8: 0.1862, 2: 0.1883, 2.4: 0.1911, 2.8: 0.1927,
              3.2: 0.1937, 3.6: 0.1943, 4: 0.1947, 5: 0.1952, 10: 0.1955},
        2.4: {1: 0.1578, 1.2: 0.1657, 1.4: 0.1715, 1.6: 0.1757, 1.8: 0.1789, 2: 0.1812, 2.4: 0.1843, 2.8: 0.1862,
              3.2: 0.1873, 3.6: 0.1873, 4: 0.1885, 5: 0.189, 10: 0.1895},
        2.6: {1: 0.1503, 1.2: 0.1583, 1.4: 0.1642, 1.6: 0.1686, 1.8: 0.1719, 2: 0.1745, 2.4: 0.1779, 2.8: 0.1799,
              3.2: 0.1812, 3.6: 0.1812, 4: 0.1825, 5: 0.1832, 10: 0.1838},
        2.8: {1: 0.1433, 1.2: 0.1514, 1.4: 0.1574, 1.6: 0.1619, 1.8: 0.1654, 2: 0.168, 2.4: 0.1717, 2.8: 0.1739,
              3.2: 0.1753, 3.6: 0.1753, 4: 0.1769, 5: 0.1777, 10: 0.1784},
        3: {1: 0.1369, 1.2: 0.1449, 1.4: 0.151, 1.6: 0.1556, 1.8: 0.1592, 2: 0.1619, 2.4: 0.1658, 2.8: 0.1682,
            3.2: 0.1698, 3.6: 0.1708, 4: 0.1715, 5: 0.1725, 10: 0.1733},
        3.2: {1: 0.131, 1.2: 0.139, 1.4: 0.145, 1.6: 0.1497, 1.8: 0.1533, 2: 0.1562, 2.4: 0.1602, 2.8: 0.1628,
              3.2: 0.1645, 3.6: 0.1657, 4: 0.1664, 5: 0.1675, 10: 0.1685},
        3.4: {1: 0.1256, 1.2: 0.1334, 1.4: 0.1394, 1.6: 0.1441, 1.8: 0.1478, 2: 0.1508, 2.4: 0.155, 2.8: 0.1577,
              3.2: 0.1595, 3.6: 0.1607, 4: 0.1616, 5: 0.1628, 10: 0.1639},
        3.6: {1: 0.1205, 1.2: 0.1282, 1.4: 0.1342, 1.6: 0.1389, 1.8: 0.1427, 2: 0.1456, 2.4: 0.15, 2.8: 0.1528,
              3.2: 0.1548, 3.6: 0.1561, 4: 0.157, 5: 0.1583, 10: 0.1595},
        3.8: {1: 0.1158, 1.2: 0.1234, 1.4: 0.1293, 1.6: 0.134, 1.8: 0.1378, 2: 0.1408, 2.4: 0.1452, 2.8: 0.1482,
              3.2: 0.1502, 3.6: 0.1516, 4: 0.1526, 5: 0.1541, 10: 0.1554},
        4: {1: 0.1114, 1.2: 0.1189, 1.4: 0.1248, 1.6: 0.1294, 1.8: 0.1332, 2: 0.1362, 2.4: 0.1408, 2.8: 0.1438,
            3.2: 0.1459, 3.6: 0.1474, 4: 0.1485, 5: 0.15, 10: 0.1516},
        4.2: {1: 0.1073, 1.2: 0.1147, 1.4: 0.1205, 1.6: 0.1251, 1.8: 0.1289, 2: 0.1319, 2.4: 0.1365, 2.8: 0.1396,
              3.2: 0.1418, 3.6: 0.1434, 4: 0.1445, 5: 0.1462, 10: 0.1479},
        4.4: {1: 0.1035, 1.2: 0.1107, 1.4: 0.1164, 1.6: 0.121, 1.8: 0.1248, 2: 0.1279, 2.4: 0.1325, 2.8: 0.1357,
              3.2: 0.1379, 3.6: 0.1396, 4: 0.1407, 5: 0.1425, 10: 0.1444},
        4.6: {1: 0.1, 1.2: 0.107, 1.4: 0.1127, 1.6: 0.1172, 1.8: 0.1209, 2: 0.124, 2.4: 0.1287, 2.8: 0.1319,
              3.2: 0.1342, 3.6: 0.1359, 4: 0.1371, 5: 0.139, 10: 0.141},
        4.8: {1: 0.0967, 1.2: 0.1036, 1.4: 0.1091, 1.6: 0.1136, 1.8: 0.1173, 2: 0.1204, 2.4: 0.125, 2.8: 0.1283,
              3.2: 0.1307, 3.6: 0.1324, 4: 0.1337, 5: 0.1357, 10: 0.1379},
        5: {1: 0.0935, 1.2: 0.1003, 1.4: 0.1057, 1.6: 0.1102, 1.8: 0.1139, 2: 0.1169, 2.4: 0.1216, 2.8: 0.1249,
            3.2: 0.1273, 3.6: 0.1291, 4: 0.1304, 5: 0.1325, 10: 0.1348},
        5.2: {1: 0.0906, 1.2: 0.0972, 1.4: 0.1026, 1.6: 0.107, 1.8: 0.1106, 2: 0.1136, 2.4: 0.1183, 2.8: 0.1217,
              3.2: 0.1241, 3.6: 0.1259, 4: 0.1273, 5: 0.1295, 10: 0.132},
        5.4: {1: 0.0878, 1.2: 0.0943, 1.4: 0.0996, 1.6: 0.1039, 1.8: 0.1075, 2: 0.1105, 2.4: 0.1152, 2.8: 0.1186,
              3.2: 0.121, 3.6: 0.1229, 4: 0.1243, 5: 0.1265, 10: 0.1292},
        5.6: {1: 0.0852, 1.2: 0.0916, 1.4: 0.0968, 1.6: 0.101, 1.8: 0.1046, 2: 0.1076, 2.4: 0.1122, 2.8: 0.1156,
              3.2: 0.1181, 3.6: 0.12, 4: 0.1215, 5: 0.1238, 10: 0.1266},
        5.8: {1: 0.0828, 1.2: 0.089, 1.4: 0.0941, 1.6: 0.0983, 1.8: 0.1018, 2: 0.1047, 2.4: 0.1094, 2.8: 0.1128,
              3.2: 0.1153, 3.6: 0.1172, 4: 0.1187, 5: 0.1211, 10: 0.124},
        6: {1: 0.0805, 1.2: 0.0866, 1.4: 0.0916, 1.6: 0.0957, 1.8: 0.0991, 2: 0.1021, 2.4: 0.1067, 2.8: 0.1101,
            3.2: 0.1126, 3.6: 0.1146, 4: 0.1161, 5: 0.1186, 10: 0.1216},
        6.2: {1: 0.0783, 1.2: 0.0842, 1.4: 0.0891, 1.6: 0.0932, 1.8: 0.0966, 2: 0.0995, 2.4: 0.1041, 2.8: 0.1075,
              3.2: 0.1101, 3.6: 0.112, 4: 0.1136, 5: 0.1161, 10: 0.1193},
        6.4: {1: 0.0762, 1.2: 0.082, 1.4: 0.0869, 1.6: 0.0909, 1.8: 0.0942, 2: 0.0971, 2.4: 0.1016, 2.8: 0.105,
              3.2: 0.1076, 3.6: 0.1096, 4: 0.1111, 5: 0.1137, 10: 0.1171},
        6.6: {1: 0.0742, 1.2: 0.0799, 1.4: 0.0847, 1.6: 0.0886, 1.8: 0.0919, 2: 0.0948, 2.4: 0.0993, 2.8: 0.1027,
              3.2: 0.1053, 3.6: 0.1073, 4: 0.1088, 5: 0.1114, 10: 0.1149},
        6.8: {1: 0.0723, 1.2: 0.0779, 1.4: 0.0826, 1.6: 0.0865, 1.8: 0.0898, 2: 0.0926, 2.4: 0.097, 2.8: 0.1004,
              3.2: 0.103, 3.6: 0.105, 4: 0.1066, 5: 0.1092, 10: 0.1129},
        7: {1: 0.0705, 1.2: 0.0761, 1.4: 0.0806, 1.6: 0.0844, 1.8: 0.0877, 2: 0.0904, 2.4: 0.0949, 2.8: 0.0982,
            3.2: 0.1008, 3.6: 0.1028, 4: 0.1044, 5: 0.1071, 10: 0.1109},
        7.2: {1: 0.0688, 1.2: 0.0742, 1.4: 0.0787, 1.6: 0.0825, 1.8: 0.0857, 2: 0.0884, 2.4: 0.0928, 2.8: 0.0962,
              3.2: 0.0987, 3.6: 0.1008, 4: 0.1023, 5: 0.1051, 10: 0.109},
        7.4: {1: 0.0672, 1.2: 0.0725, 1.4: 0.0769, 1.6: 0.0806, 1.8: 0.0838, 2: 0.0865, 2.4: 0.0908, 2.8: 0.0942,
              3.2: 0.0967, 3.6: 0.0988, 4: 0.1004, 5: 0.1031, 10: 0.1071},
        7.6: {1: 0.0656, 1.2: 0.0709, 1.4: 0.0752, 1.6: 0.0789, 1.8: 0.082, 2: 0.0846, 2.4: 0.0889, 2.8: 0.0922,
              3.2: 0.0948, 3.6: 0.0968, 4: 0.0984, 5: 0.1012, 10: 0.1054},
        7.8: {1: 0.0642, 1.2: 0.0693, 1.4: 0.0736, 1.6: 0.0771, 1.8: 0.0802, 2: 0.0828, 2.4: 0.0871, 2.8: 0.0904,
              3.2: 0.0929, 3.6: 0.095, 4: 0.0966, 5: 0.0994, 10: 0.1036},
        8: {1: 0.0627, 1.2: 0.0678, 1.4: 0.072, 1.6: 0.0755, 1.8: 0.0785, 2: 0.0811, 2.4: 0.0853, 2.8: 0.0886,
            3.2: 0.0912, 3.6: 0.0932, 4: 0.0948, 5: 0.0976, 10: 0.102},
        8.2: {1: 0.0614, 1.2: 0.0663, 1.4: 0.0705, 1.6: 0.0739, 1.8: 0.0769, 2: 0.0795, 2.4: 0.0837, 2.8: 0.0869,
              3.2: 0.0894, 3.6: 0.0914, 4: 0.0931, 5: 0.0959, 10: 0.1004},
        8.4: {1: 0.0601, 1.2: 0.0649, 1.4: 0.069, 1.6: 0.0724, 1.8: 0.0754, 2: 0.0779, 2.4: 0.082, 2.8: 0.0852,
              3.2: 0.0878, 3.6: 0.0893, 4: 0.0914, 5: 0.0943, 10: 0.0938},
        8.6: {1: 0.0588, 1.2: 0.0636, 1.4: 0.0676, 1.6: 0.071, 1.8: 0.0739, 2: 0.0764, 2.4: 0.0805, 2.8: 0.0836,
              3.2: 0.0862, 3.6: 0.0882, 4: 0.0898, 5: 0.0927, 10: 0.0973},
        8.8: {1: 0.0576, 1.2: 0.0623, 1.4: 0.0663, 1.6: 0.0606, 1.8: 0.0724, 2: 0.0749, 2.4: 0.079, 2.8: 0.0821,
              3.2: 0.0846, 3.6: 0.0866, 4: 0.0882, 5: 0.0912, 10: 0.0959},
        9.2: {1: 0.0554, 1.2: 0.0559, 1.4: 0.0637, 1.6: 0.067, 1.8: 0.0697, 2: 0.0721, 2.4: 0.0761, 2.8: 0.0792,
              3.2: 0.0817, 3.6: 0.0837, 4: 0.0853, 5: 0.0882, 10: 0.0931},
        9.6: {1: 0.0533, 1.2: 0.0577, 1.4: 0.0614, 1.6: 0.0645, 1.8: 0.0672, 2: 0.0696, 2.4: 0.0734, 2.8: 0.0765,
              3.2: 0.0789, 3.6: 0.0809, 4: 0.0825, 5: 0.0855, 10: 0.0905},
        10: {1: 0.0514, 1.2: 0.0556, 1.4: 0.0592, 1.6: 0.0622, 1.8: 0.0649, 2: 0.0672, 2.4: 0.071, 2.8: 0.0739,
             3.2: 0.0763, 3.6: 0.0783, 4: 0.0799, 5: 0.0829, 10: 0.088},
        10.4: {1: 0.0496, 1.2: 0.0537, 1.4: 0.0572, 1.6: 0.0601, 1.8: 0.0627, 2: 0.0649, 2.4: 0.0686, 2.8: 0.0716,
               3.2: 0.0739, 3.6: 0.0759, 4: 0.0775, 5: 0.0804, 10: 0.0857},
        10.8: {1: 0.0479, 1.2: 0.0519, 1.4: 0.0553, 1.6: 0.0581, 1.8: 0.0606, 2: 0.0628, 2.4: 0.0664, 2.8: 0.0693,
               3.2: 0.0717, 3.6: 0.0736, 4: 0.0751, 5: 0.0781, 10: 0.0834},
        11.2: {1: 0.0463, 1.2: 0.0502, 1.4: 0.0502, 1.6: 0.0563, 1.8: 0.0587, 2: 0.0609, 2.4: 0.0664, 2.8: 0.0672,
               3.2: 0.0695, 3.6: 0.0714, 4: 0.073, 5: 0.0759, 10: 0.0813},
        11.6: {1: 0.0448, 1.2: 0.0486, 1.4: 0.0518, 1.6: 0.0545, 1.8: 0.0569, 2: 0.059, 2.4: 0.0625, 2.8: 0.0652,
               3.2: 0.0675, 3.6: 0.0694, 4: 0.0709, 5: 0.0738, 10: 0.0793},
        12: {1: 0.0435, 1.2: 0.0471, 1.4: 0.0502, 1.6: 0.0529, 1.8: 0.0552, 2: 0.0573, 2.4: 0.0606, 2.8: 0.0634,
             3.2: 0.0656, 3.6: 0.0674, 4: 0.069, 5: 0.0719, 10: 0.0774},
        12.8: {1: 0.0409, 1.2: 0.0444, 1.4: 0.0474, 1.6: 0.0499, 1.8: 0.0521, 2: 0.0541, 2.4: 0.0573, 2.8: 0.0599,
               3.2: 0.0621, 3.6: 0.0639, 4: 0.0654, 5: 0.0682, 10: 0.0739},
        13.6: {1: 0.0387, 1.2: 0.042, 1.4: 0.0448, 1.6: 0.0472, 1.8: 0.0493, 2: 0.0512, 2.4: 0.0543, 2.8: 0.0568,
               3.2: 0.0589, 3.6: 0.0607, 4: 0.0621, 5: 0.0649, 10: 0.0707},
        14.4: {1: 0.0367, 1.2: 0.0398, 1.4: 0.0425, 1.6: 0.0488, 1.8: 0.0468, 2: 0.0486, 2.4: 0.0516, 2.8: 0.054,
               3.2: 0.0561, 3.6: 0.0577, 4: 0.0592, 5: 0.0619, 10: 0.0677},
        15.2: {1: 0.0349, 1.2: 0.0379, 1.4: 0.0404, 1.6: 0.0426, 1.8: 0.0446, 2: 0.0463, 2.4: 0.0492, 2.8: 0.0515,
               3.2: 0.0535, 3.6: 0.0551, 4: 0.0565, 5: 0.0592, 10: 0.065},
        16: {1: 0.0332, 1.2: 0.0361, 1.4: 0.0385, 1.6: 0.0407, 1.8: 0.0425, 2: 0.0442, 2.4: 0.0469, 2.8: 0.0492,
             3.2: 0.0511, 3.6: 0.0527, 4: 0.054, 5: 0.0567, 10: 0.0625},
        18: {1: 0.0297, 1.2: 0.0323, 1.4: 0.0345, 1.6: 0.0364, 1.8: 0.0381, 2: 0.0396, 2.4: 0.0422, 2.8: 0.0442,
             3.2: 0.046, 3.6: 0.0475, 4: 0.0487, 5: 0.0512, 10: 0.057},
        20: {1: 0.0269, 1.2: 0.0292, 1.4: 0.0312, 1.6: 0.033, 1.8: 0.0345, 2: 0.0359, 2.4: 0.0383, 2.8: 0.0402,
             3.2: 0.0418, 3.6: 0.0432, 4: 0.0444, 5: 0.0468, 10: 0.0524}}
    z_b_values = [0, 0.2, 0.4, 0.6, 0.8, 1, 1.2, 1.4, 1.6, 1.8, 2, 2.2, 2.4, 2.6, 2.8, 3, 3.2, 3.4, 3.6, 3.8, 4, 4.2,
                  4.4, 4.6, 4.8, 5, 5.2, 5.4, 5.6, 5.8, 6, 6.2, 6.4, 6.6, 6.8, 7, 7.2, 7.4, 7.6, 7.8, 8, 8.2, 8.4, 8.6,
                  8.8, 9.2, 9.6, 10, 10.4, 10.8, 11.2, 11.6, 12, 12.8, 13.6, 14.4, 15.2, 16, 18, 20]
    l_b_values = [1, 1.2, 1.4, 1.6, 1.8, 2, 2.4, 2.8, 3.2, 3.6, 4, 5, 10]
    return alpha_data, z_b_values, l_b_values


def interpolate_alpha(alpha_data, z_b_values, l_b_values, z_b, l_b):
    # Find closest z_b
    z_b_sorted = sorted(z_b_values)
    l_b_sorted = sorted(l_b_values)

    # Find indices for z_b
    if z_b <= z_b_sorted[0]:
        z1 = z2 = z_b_sorted[0]
    elif z_b >= z_b_sorted[-1]:
        z1 = z2 = z_b_sorted[-1]
    else:
        for i in range(len(z_b_sorted) - 1):
            if z_b_sorted[i] <= z_b <= z_b_sorted[i + 1]:
                z1 = z_b_sorted[i]
                z2 = z_b_sorted[i + 1]
                break

    # Find indices for l_b
    if l_b <= l_b_sorted[0]:
        l1 = l2 = l_b_sorted[0]
    elif l_b >= l_b_sorted[-1]:
        l1 = l2 = l_b_sorted[-1]
    else:
        for i in range(len(l_b_sorted) - 1):
            if l_b_sorted[i] <= l_b <= l_b_sorted[i + 1]:
                l1 = l_b_sorted[i]
                l2 = l_b_sorted[i + 1]
                break

    if z1 == z2 and l1 == l2:
        return alpha_data[z1][l1]

    # Get the four points
    a11 = alpha_data[z1].get(l1, 0)
    a12 = alpha_data[z1].get(l2, 0)
    a21 = alpha_data[z2].get(l1, 0)
    a22 = alpha_data[z2].get(l2, 0)

    # Bilinear interpolation
    if z1 != z2 and l1 != l2:
        r1 = a11 + (a12 - a11) * (l_b - l1) / (l2 - l1)
        r2 = a21 + (a22 - a21) * (l_b - l1) / (l2 - l1)
        return r1 + (r2 - r1) * (z_b - z1) / (z2 - z1)
    elif z1 != z2:
        return a11 + (a21 - a11) * (z_b - z1) / (z2 - z1)
    elif l1 != l2:
        return a11 + (a12 - a11) * (l_b - l1) / (l2 - l1)
    else:
        return a11


def add_deformation_section(doc, buildings, buildings_list, holes, hole_strata, layer_info, building_results):
    deformation_buildings = []

    # 1. 筛选需要计算的建筑物（所有高层建筑，即高度>24m）
    for building in buildings_list:
        h_raw = buildings.get(building, {}).get('height')
        try:
            height = float(str(h_raw).strip()) if h_raw not in (None, '', '/', '-', 'None', 'null') else 0.0
        except (ValueError, TypeError):
            height = 0.0
            logging.warning(f"建筑物 {building} 的高度 {h_raw} 无效，设为 0")

        # 只要高度大于24m，全部纳入计算
        if height <= 24:
            continue
        deformation_buildings.append(building)

    if not deformation_buildings:
        logging.info("无高层建筑物需要变形估算，未生成表6")
        return

    alpha_data, z_b_values, l_b_values = get_alpha_data()

    doc.add_heading('（2）地基变形评价', level=3)
    doc.add_paragraph(
        f"根据《高层建筑岩土工程勘察标准》(JGJ/T 72-2017)的技术要求，对拟建高层建筑"
        f"{'、'.join(deformation_buildings)}进行地基变形估算，估算结果见表6。"
    )
    doc.add_paragraph("                                                               地基变形估算    表6")
    table6 = doc.add_table(rows=1, cols=6)
    table6.style = 'Table Grid'

    hdr = table6.rows[0].cells
    hdr[0].text = '拟建建筑'
    hdr[1].text = '项目'
    hdr[2].text = '左下角点'
    hdr[3].text = '右下角点'
    hdr[4].text = '右上角点'
    hdr[5].text = '左上角点'

    def get_real_corners(build_holes_dict):
        points = [(info.get('x'), info.get('y'), h)
                  for h, info in build_holes_dict.items()
                  if info.get('x') is not None and info.get('y') is not None]
        if len(points) < 3:
            return ['', '', '', '']
        xs = [p[0] for p in points]
        ys = [p[1] for p in points]
        min_x, max_x = min(xs), max(xs)
        min_y, max_y = min(ys), max(ys)
        corners = [(min_x, min_y), (max_x, min_y), (max_x, max_y), (min_x, max_y)]
        used = set()
        result = []
        for tx, ty in corners:
            candidates = [p for p in points if p[2] not in used]
            if not candidates:
                result.append('')
                continue
            nearest = min(candidates, key=lambda p: (p[0] - tx) ** 2 + (p[1] - ty) ** 2)
            result.append(nearest[2])
            used.add(nearest[2])
        return result

    # === 修改点1：修正倾斜限值逻辑 ===
    def safe_tilt_limit(height_val):
        if height_val is None or str(height_val).strip() in {'', '-', '/', '—', '－', '无', '暂无', 'None', 'null'}:
            return 0.004
        try:
            h = float(str(height_val).strip())
            if h <= 0: return 0.004

            # 按新规则判断
            if h <= 24:
                return 0.004
            elif h <= 60:
                return 0.003
            elif h <= 100:
                return 0.0025
            else:
                return 0.002
        except (ValueError, TypeError):
            return 0.004

    summary = {}

    for building in deformation_buildings:
        build_holes = {h: info for h, info in holes.items()
                       if
                       building in info.get('builds', []) and info.get('x') is not None and info.get('y') is not None}

        if len(build_holes) < 3:
            logging.warning(f"{building} 坐标点少于3个，跳过")
            building_results[building]['deform_satisfy'] = False
            continue

        hole_lb, hole_rb, hole_rt, hole_lt = get_real_corners(build_holes)

        def safe_s(hole):
            return calculate_settlement(building, hole, buildings, holes, hole_strata, layer_info,
                                        alpha_data, z_b_values, l_b_values) if hole else 0.0

        s_lb = safe_s(hole_lb)
        s_rb = safe_s(hole_rb)
        s_rt = safe_s(hole_rt)
        s_lt = safe_s(hole_lt)

        diff_x = (abs(s_lb - s_rb) + abs(s_lt - s_rt)) / 2
        diff_y = (abs(s_lb - s_lt) + abs(s_rb - s_rt)) / 2

        B = buildings[building].get('width') or buildings[building].get('B')
        L = buildings[building].get('length') or buildings[building].get('L')

        if not B or B <= 0:
            xs = [info['x'] for info in build_holes.values() if info['x'] is not None]
            B = max(xs) - min(xs) if len(xs) >= 2 else 30
        if not L or L <= 0:
            ys = [info['y'] for info in build_holes.values() if info['y'] is not None]
            L = max(ys) - min(ys) if len(ys) >= 2 else 50

        tilt_x = diff_x / L / 1000 if L > 0 else 0
        tilt_y = diff_y / B / 1000 if B > 0 else 0

        # 获取该建筑的倾斜限值
        limit = safe_tilt_limit(buildings[building].get('height'))

        # === 修改点2：增加沉降量判断（需同时满足倾斜和沉降<=200mm）===
        max_settlement = max(s_lb, s_rb, s_rt, s_lt)

        # 判定条件：X向倾斜达标 AND Y向倾斜达标 AND 最大沉降量 <= 200
        deform_ok = (tilt_x <= limit) and (tilt_y <= limit) and (max_settlement <= 200)

        building_results[building]['deform_satisfy'] = deform_ok

        base_layers, _, _, _ = get_base_layers(building, buildings, holes, hole_strata, layer_info)
        summary[building] = {'ok': deform_ok, 'layers': base_layers or ['相应持力层']}

        # 填表
        row = table6.add_row().cells
        row[0].text = building
        row[1].text = '角点孔号'
        row[2].text = hole_lb or '-'
        row[3].text = hole_rb or '-'
        row[4].text = hole_rt or '-'
        row[5].text = hole_lt or '-'

        row = table6.add_row().cells
        row[1].text = '沉降量(mm)'
        row[2].text = f"{s_lb:.1f}" if hole_lb else '-'
        row[3].text = f"{s_rb:.1f}" if hole_rb else '-'
        row[4].text = f"{s_rt:.1f}" if hole_rt else '-'
        row[5].text = f"{s_lt:.1f}" if hole_lt else '-'

        row = table6.add_row().cells
        row[1].text = '沉降差(mm)'
        row[2].text = f"{diff_x:.1f}"
        row[4].text = f"{diff_y:.1f}"

        row = table6.add_row().cells
        row[1].text = '倾斜'
        row[2].text = f"{tilt_x:.4f}"
        row[4].text = f"{tilt_y:.4f}"

        row = table6.add_row().cells
        row[1].text = '评价'

        # 评价列显示逻辑：如果不满足，需显示具体原因
        if deform_ok:
            row[2].text = "满足"
            row[4].text = "满足"
        else:
            reason = []
            if max_settlement > 200: reason.append("沉降过大")
            if tilt_x > limit or tilt_y > limit: reason.append("倾斜过大")
            fail_text = "不满足" + ("(" + ",".join(reason) + ")" if reason else "")
            row[2].text = fail_text
            row[4].text = fail_text

    doc.add_paragraph(
        "说明：①表中沉降量为最终沉降估算值；②倾斜值为X、Y向平均倾斜；③本表为初步估算，设计单位应根据结构实际荷载、刚度分布等进行详细计算与验算；④根据规范要求，沉降量应不大于200mm。")

    ok = [b for b, v in summary.items() if v['ok']]
    nok = [b for b, v in summary.items() if not v['ok']]

    if ok:
        layers = set()
        for b in ok:
            layers.update(summary[b]['layers'])
        layers_str = '、'.join(sorted(list(set([l.split('（')[0] for l in layers]))))
        doc.add_paragraph(f"根据表6估算结果，{'、'.join(ok)}地基变形满足规范要求。")
        doc.add_paragraph(
            f"对{'、'.join(ok)}基底分布地层为较好的{layers_str}，结合大量工程经验，若其采用天然地基独立基础或筏形基础，地基变形一般能满足要求。")

    if nok:
        doc.add_paragraph(f"根据表6估算结果，{'、'.join(nok)}地基变形不满足规范要求，从变形角度考虑不宜采用天然地基。")


# --- 主分析函数 ---
def parse_hole_buildings_from_sheet1_5(wb):
    """从 1.5单孔 表的 M列（建筑物名称）建立孔号 → 楼栋列表"""
    sheet = wb['1.5单孔']
    hole_to_buildings = {}

    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not row or not row[0]:  # A列孔号为空，跳过
            continue
        hole_id = str(row[0]).strip()
        if len(row) <= 12:  # M列是第13列（索引12）
            continue
        building_names = row[12]
        if not building_names:
            continue

        # 支持多种分隔符：, 、 ; 空格
        names = re.split(r'[,，;；\s]+', str(building_names).strip())
        clean_names = [n.strip() for n in names if n.strip()]

        if clean_names:
            hole_to_buildings[hole_id] = clean_names
            print(f"孔 {hole_id} 关联楼栋: {clean_names}")  # 调试用

    return hole_to_buildings




def run_analysis(template_path, input_path, output_path, ask_yes_no):
    logging.info(f"开始分析: 输入={input_path}, 输出={output_path}")
    try:
        wb = openpyxl.load_workbook(input_path, data_only=True)
        # 打开模板或新建
        doc = Document(template_path or None)

        # ====================== 关键：立即设置样式 ======================
        set_global_styles(doc)
        force_set_all_fonts(doc)
        # =============================================================
        # 解析数据
        buildings, buildings_list = parse_buildings(wb['1.1'])
        holes = parse_holes(wb['1.5单孔'], buildings_list)

        layer_info = parse_layer_info(wb['1.6地层信息'])

        def safe_float(val, default=0.0):
            try:
                return float(val) if val is not None else default
            except:
                return default

        # 把所有可能为空的数值字段强制转成 float，None → 0
        for lid in layer_info:
            info = layer_info[lid]
            info['bearing_capacity'] = safe_float(info.get('bearing_capacity'), 0)
            info['compression_modulus'] = safe_float(info.get('compression_modulus'), 0)
            info['density'] = safe_float(info.get('density'), 0)
        # ============================================================
        hole_strata = parse_hole_strata(wb['2.4各孔地层'], holes)

        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()

        # 初始化 building_results，包含所有必要的默认值
        building_results = defaultdict(lambda: {
            'ind_unsatisfy': [],
            'raft_unsatisfy': [],
            'all_above_ground': False,
            'avg_above_height': 0.0,
            'above_ground_info': [],
            'base_layers': [],
            # === 新增状态标志，默认为 False，后续根据计算更新 ===
            'ind_cap_ok': False,  # 独立基础持力层承载力
            'raft_cap_ok': False, # 筏板基础持力层承载力
            'ind_weak_ok': True,  # 独立基础软弱下卧层 (默认True，无下卧层即为满足)
            'raft_weak_ok': True, # 筏板基础软弱下卧层 (默认True)
            'deform_satisfy': True # 变形 (非高层默认True)
        })

        # --- 第一节：地基均匀性评价 ---
        doc.add_heading('地基基础分析评价及选型', level=1)
        doc.add_heading('1、地基均匀性评价', level=2)
        doc.add_paragraph('根据勘察揭露地层分布情况，结合拟建建筑性质，各拟建建筑基底标高处分布地层情况见表1。')

        # 生成表1：基底分布地层情况
        doc.add_paragraph('\n                                                                       基底分布地层情况  表1')
        table1 = doc.add_table(rows=1, cols=2)
        table1.style = 'Table Grid'

        # === 【新增】设置列宽 ===
        table1.autofit = False  # 关键：关闭自动适应，否则宽度设置无效

        # 设置第一列宽度为 2cm
        table1.columns[0].width = Cm(2)

        # 建议：设置第二列宽度（A4纸除去页边距约剩16cm，这里设为14cm填满剩余空间）
        # 如果不设置第二列，Word有时会显示异常，建议显式指定
        table1.columns[1].width = Cm(14.5)
        # =======================

        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = '拟建建筑'
        hdr_cells[1].text = '基底分布地层情况'

        # 双重保险：显式设置表头单元格宽度（防止Word某些版本不认列属性）
        hdr_cells[0].width = Cm(2)
        hdr_cells[1].width = Cm(14.5)

        fill_buildings = []
        uniform_buildings = []
        for building in buildings_list:
            desc = compute_desc(building, buildings, holes, layer_info, hole_strata)
            row_cells = table1.add_row().cells
            row_cells[0].text = building
            row_cells[1].text = desc
            logging.info(f"建筑物 {building} 的描述: {desc}")

            # 判断是否为填土或单一地层
            build_holes_list = [h for h, info in holes.items() if building in info['builds']]
            hole_to_layer = {}
            for hole_id in build_holes_list:
                hole_elev = holes[hole_id]['elev']
                if hole_elev is None:
                    continue
                embed_depth = hole_elev - buildings[building]['embed_elev']
                if embed_depth < 0:
                    continue
                layer = find_layer_at_depth(hole_strata, holes, hole_id, embed_depth)
                if layer:
                    hole_to_layer[hole_id] = layer
            layer_counts = Counter(hole_to_layer.values())
            if any(is_fill(layer) for layer in layer_counts):
                fill_buildings.append(building)
            elif len(layer_counts) == 1 and all(
                    holes[h]['elev'] is not None and holes[h]['elev'] >= buildings[building]['embed_elev'] for h in
                    build_holes_list):
                uniform_buildings.append((building, layer_counts.most_common(1)[0][0]))

        doc.add_paragraph('\n各拟建建筑基底以下地层分布具体情况见表2')
        doc.add_paragraph('                                                               基底各钻孔分布具体地层  表2')
        table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
        # === 【新增】设置列宽 ===
        table2.autofit = False  # 关键：关闭自动适应，否则宽度设置无效

        # 设置第一列宽度为 2cm
        table2.columns[0].width = Cm(2)

        # 建议：设置第二列宽度（A4纸除去页边距约剩16cm，这里设为14cm填满剩余空间）
        # 如果不设置第二列，Word有时会显示异常，建议显式指定
        table2.columns[1].width = Cm(2)
        table2.columns[2].width = Cm(12.5)
        # =======================

        h = table2.rows[0].cells
        h[0].text, h[1].text, h[2].text = '楼栋号', '钻孔', '基底以下层号/深度'

        # 填充表2
        for building in buildings_list:
            build_holes_list = [h for h, info in holes.items() if building in info.get('builds', [])]
            if not build_holes_list:
                row_cells = table2.add_row().cells
                row_cells[0].text = building
                row_cells[1].text, row_cells[2].text = '', '无勘探点数据'
                continue

            base_elev = buildings[building]['embed_elev']
            for hole_id in build_holes_list:
                row_cells = table2.add_row().cells
                row_cells[0].text, row_cells[1].text = building, hole_id

                hole_data = holes.get(hole_id, {})
                hole_elev = hole_data.get('elev')
                if hole_elev is None:
                    row_cells[2].text = '无标高数据'
                    continue

                embed_depth = hole_elev - base_elev
                strata = hole_strata.get(hole_id, [])
                if not strata:
                    row_cells[2].text = '无地层数据'
                    continue

                layers_desc_list, prev_depth = [], 0.0
                if embed_depth < 0:
                    layers_desc_list.append(f"基底高于地面约{-embed_depth:.2f}m")
                    for layer_id, bottom in strata:
                        name = layer_info.get(layer_id, {}).get('name', '')
                        layers_desc_list.append(f"{name}{layer_id} ({prev_depth:.2f}-{bottom:.2f}m)")
                        prev_depth = bottom
                else:
                    found_base = False
                    for layer_id, bottom in strata:
                        if not found_base and embed_depth <= bottom: found_base = True
                        if found_base:
                            start_depth = max(embed_depth, prev_depth)
                            if start_depth < bottom:
                                name = layer_info.get(layer_id, {}).get('name', '')
                                layers_desc_list.append(f"{name}{layer_id} ({start_depth:.2f}-{bottom:.2f}m)")
                        prev_depth = bottom
                row_cells[2].text = '，'.join(layers_desc_list) if layers_desc_list else '无基底以下地层'

        # 添加均匀性描述
        # 【步骤1：统计哪些楼基底在填土上，哪些楼基底在同一天然土层（均匀）】
        fill_buildings = []  # 基底落在填土上的楼栋名
        uniform_buildings = {}  # 基底全在同一个天然土层上的楼栋 → {building: layer_id}

        for building in buildings_list:
            # 获取该楼关联的所有钻孔
            assoc_holes = [h for h, info in holes.items() if building in info.get('builds', [])]
            if not assoc_holes:
                continue

            base_layers = set()  # 这栋楼所有钻孔在基底深度处的地层ID（去重）
            has_fill = False  # 这栋楼是否至少有一个孔基底落在填土层

            for hole_id in assoc_holes:
                # 计算该钻孔相对于这栋楼基础埋深标高的嵌入深度
                hole_elev = holes[hole_id].get('elev')
                embed_elev = buildings[building]['embed_elev']
                if hole_elev is None or embed_elev is None:
                    continue
                embed_depth = hole_elev - embed_elev
                if embed_depth <= 0:
                    continue  # 基底高于地面，不参与判断

                # 查找基底标高所在的地层
                layer_id = None
                prev_bottom = 0.0
                for lid, bottom in hole_strata.get(hole_id, []):
                    if prev_bottom <= embed_depth <= bottom:
                        layer_id = lid
                        break
                    prev_bottom = bottom
                else:
                    # 没找到就取最后一层（兜底）
                    if hole_strata.get(hole_id):
                        layer_id = hole_strata[hole_id][-1][0]

                if layer_id:
                    if is_fill(layer_id):
                        has_fill = True
                    else:
                        base_layers.add(layer_id)

            # 判断这栋楼归类到哪里
            if has_fill:
                fill_buildings.append(building)
            elif len(base_layers) == 1:  # 所有非填土孔都在同一个天然土层 → 均匀
                main_layer_id = next(iter(base_layers))
                uniform_buildings[building] = main_layer_id

        # 【步骤2：生成最终要写进报告的文字】
        desc_parts = []

        # 有填土的楼 → 说不均匀
        if fill_buildings:
            if len(fill_buildings) == len(buildings_list):
                desc_parts.append(
                    "根据表1可知，拟建所有建筑物基底均分布有填土，填土厚度及性质不均匀，可初步判断为不均匀地基。")
            else:
                # 直接连接所有建筑名称，不再切片取前3个
                all_fill_str = '、'.join(fill_buildings)
                desc_parts.append(
                    f"根据表1可知，拟建的{all_fill_str}基底分布有填土，填土厚度及性质不均匀，可初步判断为不均匀地基。")

        # 有均匀的楼 → 追加说明
        if uniform_buildings:
            uniform_texts = []
            for building, layer_id in uniform_buildings.items():
                info = layer_info.get(layer_id, {})
                layer_name = info.get('name', '').strip()
                state = info.get('state', '').strip()
                full_name = f"{layer_name}{state}{layer_id}" if state and state != '/' else f"{layer_name}{layer_id}"
                uniform_texts.append(f"{building}基底分布为{full_name}，为均匀地基")

            # 如果前面有话，加分号连接；否则直接开始
            sep = "；" if desc_parts else ""
            desc_parts.append(sep + "；".join(uniform_texts) + "。")

        if not desc_parts:
            desc_parts.append("根据表1可知，拟建建筑物基底持力层分布较均匀，属较均匀地基。")

        doc.add_paragraph("".join(desc_parts))

        # 处理高层建筑的当量模量评价
        modulus_buildings_info = []  # 存储 (building_name, mod_data, is_replaced)

        for b in buildings_list:
            # 1. 判断是否需要计算，以及是否涉及换填
            needs_calc, is_replaced = needs_equivalent_modulus(b, holes, hole_strata, buildings)

            if needs_calc:
                # 2. 执行计算 (传入 is_replace_mode 参数)
                mod_data = compute_equivalent_modulus(building=b, holes=holes, hole_strata=hole_strata,
                                                      buildings=buildings, layer_info=layer_info,
                                                      is_replace_mode=is_replaced)
                if mod_data:
                    modulus_buildings_info.append((b, mod_data, is_replaced))

        if modulus_buildings_info:
            doc.add_paragraph('\n对部分高层建筑进行当量模量均匀性评价，见表3。')
            doc.add_paragraph('                                                                       地基土均匀性评价  表3')
            table3 = doc.add_table(rows=1, cols=7, style='Table Grid')
            h = table3.rows[0].cells
            h[0].text, h[1].text, h[2].text = '拟建建筑', 'Es,max(MPa)', 'Es,min(MPa)'
            h[3].text, h[4].text, h[5].text = 'Es,max/Es,min', 'Es,avg(MPa)', '界限值 K'
            h[6].text = '均匀性判定'

            for b_name, mod_data, is_replaced in modulus_buildings_info:
                row = table3.add_row().cells
                row[0].text = b_name
                row[1].text = str(mod_data['es_max'])
                row[2].text = str(mod_data['es_min'])
                row[3].text = str(mod_data['es_ratio'])
                row[4].text = str(mod_data['es_avg'])
                row[5].text = str(mod_data['k'])
                row[6].text = mod_data['uniformity']

            # 生成总结文字
            summary_parts = []
            for b_name, mod_data, is_replaced in modulus_buildings_info:
                uni = mod_data['uniformity']
                if is_replaced:
                    summary_parts.append(
                        f"{b_name}为{uni}地基（该楼栋存在厚度可换填填土，均匀性计算为考虑换填以后的计算结果）")
                else:
                    summary_parts.append(f"{b_name}为{uni}地基")

            summary = "根据表3可知，" + '、'.join(summary_parts) + "。"
            doc.add_paragraph(summary)

        # --- 第二节：天然地基评价 ---
        doc.add_heading('\n2、天然地基评价', level=2)
        doc.add_heading('（1）地基承载力评价', level=3)

        is_above = ask_yes_no("持力层位置", "持力层位于地下水位以上吗？（是: γ=20, 否: γ=10）")
        gamma = 20 if is_above else 10
        gamma_m = gamma
        logging.info(f"选择γ和γm为 {gamma}")
        silt_params, clay_params = {}, {}
        all_lnames = {info['name'] for info in layer_info.values()}
        if any('粉土' in n for n in all_lnames):
            silt_params['粉土'] = (0.3, 1.5) if ask_yes_no("粉土", "粉土黏粒含量是否≥10%?") else (0.5, 2.0)
        if any('粉质黏土' in n or '粉质粘土' in n for n in all_lnames):
            clay_params['粉质黏土'] = (0.0, 1.0) if ask_yes_no("粉质黏土", "粉质黏土e或IL≥0.85?") else (0.3,
                                                                                                                 1.6)

        # 收集所有独特基底地层
        all_base_layers = set()
        for building in buildings_list:
            base_layers, _, _, _ = get_base_layers(building, buildings, holes, hole_strata, layer_info)
            all_base_layers.update(base_layers)
        base_layers_str = '、'.join(sorted(all_base_layers))
        has_clay = any(
            '黏土' in layer or '粘土' in layer or '粉质黏土' in layer or '粉质粘土' in layer or '膨胀土' in layer for
            layer in all_base_layers)
        norms = f"根据《建筑地基基础设计规范》(GB50007-2011){'及《膨胀土地区建筑技术规范》(GB50112-2013)' if has_clay else ''}规定及预估的拟建建筑基础情况，对基底下分布地层的承载力特征值进行修正估算。修正结果见表4。"
        doc.add_paragraph(norms)
        doc.add_paragraph('                                                              地基承载力估算  表4')
        table4 = doc.add_table(rows=1, cols=6, style='Table Grid')
        h = table4.rows[0].cells
        h[0].text, h[1].text, h[2].text = '拟建建筑', '基底地层', '基础形式'
        h[3].text, h[4].text, h[5].text = 'fa(kPa)', 'pk(kPa)', '是否满足荷载要求'

        # 处理每个建筑物
        for building in buildings_list:
            base_layers, all_above_ground, avg_above_height, above_ground_info = get_base_layers(
                building, buildings, holes, hole_strata, layer_info
            )

            # 存储到 building_results
            building_results[building].update({
                'base_layers': base_layers,
                'all_above_ground': all_above_ground,
                'avg_above_height': avg_above_height,
                'above_ground_info': above_ground_info
            })

            load_kpa = buildings[building].get('load')
            load_str = f"{load_kpa:.2f}" if load_kpa is not None else '/'

            # 临时列表用于判断该楼栋整体情况
            current_ind_unsat = []
            current_raft_unsat = []

            # --- 情况A: 基底高于地面 ---
            if all_above_ground and above_ground_info:
                for layer_str, ind_satisfy, raft_satisfy in above_ground_info:
                    r1 = table4.add_row().cells
                    r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text, r1[5].text = (
                        building, layer_str, '独立基础', '/', load_str, ind_satisfy)
                    r2 = table4.add_row().cells
                    r2[0].text, r2[2].text, r2[3].text, r2[4].text, r2[5].text = (
                        '', '筏形基础', '/', load_str, raft_satisfy)

                    if ind_satisfy == '不满足': current_ind_unsat.append(layer_str)
                    if raft_satisfy == '不满足': current_raft_unsat.append(layer_str)

            # --- 情况B: 正常地下基底 ---
            else:
                for layer_str in base_layers:
                    orig_lstr = layer_str.split('（')[0]
                    l_id = next((lid for lid, i in layer_info.items() if f"{i.get('name', '')}{lid}" == orig_lstr),
                                None)
                    if not l_id: continue

                    l_name, fak = layer_info[l_id]['name'], layer_info[l_id]['bearing_capacity']
                    eta_b, eta_d, is_clay = get_eta_params(l_name, silt_params, clay_params)

                    # 独立基础计算
                    fa_ind = calculate_fa(fak, eta_b, eta_d, is_clay, 3, 1.5, l_name, gamma, gamma)
                    s_ind = '满足' if isinstance(fa_ind, (int,
                                                          float)) and load_kpa is not None and fa_ind >= load_kpa else '不满足'
                    if s_ind == '不满足': current_ind_unsat.append(layer_str)

                    r1 = table4.add_row().cells
                    r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text, r1[5].text = (
                        building, layer_str, '独立基础', str(fa_ind), load_str, s_ind)

                    # 筏板基础计算
                    fa_raft = calculate_fa(fak, eta_b, eta_d, is_clay, 6, 1.5, l_name, gamma, gamma)
                    s_raft = '满足' if isinstance(fa_raft, (int,
                                                            float)) and load_kpa is not None and fa_raft >= load_kpa else '不满足'
                    if s_raft == '不满足': current_raft_unsat.append(layer_str)

                    r2 = table4.add_row().cells
                    r2[0].text, r2[2].text, r2[3].text, r2[4].text, r2[5].text = (
                        '', '筏形基础', str(fa_raft), load_str, s_raft)

            # 更新全局状态
            building_results[building]['ind_unsatisfy'] = current_ind_unsat
            building_results[building]['raft_unsatisfy'] = current_raft_unsat
            # 关键：设置承载力是否满足的标志位
            building_results[building]['ind_cap_ok'] = (len(current_ind_unsat) == 0)
            building_results[building]['raft_cap_ok'] = (len(current_raft_unsat) == 0)

            # 添加表4说明（保持不变）
        doc.add_paragraph(f'说明：①持力层按{"地下水位以上" if is_above else "地下水位以下"}考虑，γ={gamma}，γm={gamma_m}。')
        doc.add_paragraph(
            '②各拟建建筑独立基础或筏形基础深度均按1.5m考虑；独立基础宽度按不大于3.0m考虑，筏形基础宽度按不小于6.0m考虑；')
        doc.add_paragraph(
            '③黏土按公式fa=fak+ηb*γ*(b-3)+ηd*γm*(d-1)进行修正，其余地层按公式fa=fak+ηb*γ*(b-3)+ηd*γm*(d-0.5)进行修正，中等风化岩不进行修正。')

        # === 修改后的表4总结逻辑 ===
        summary_text = "根据表4估算结果，结合各拟建建筑基底分布地层情况表可知：\n"

        for building in buildings_list:
            res = building_results[building]
            if res['all_above_ground']:
                # 基底高于地面的特殊处理
                avg_height = building_results[building]['avg_above_height']
                above_info = building_results[building]['above_ground_info']
                layers_str = '、'.join([layer for layer, _, _ in above_info])
                main_layer, ind_satisfy, raft_satisfy = above_info[0]
                summary_text += f"{building}基底高于地面{avg_height:.2f}m，地表主要分布地层{layers_str}，"
                if res['ind_cap_ok'] and res['raft_cap_ok']:
                    summary_text += "采用独立基础或筏形基础均满足承载力要求。\n"
                elif res['raft_cap_ok']:
                    summary_text += "采用独立基础不满足承载力要求，采用筏形基础满足承载力要求。\n"
                else:
                    summary_text += "采用独立基础或筏形基础均不满足承载力要求。\n"
                continue

            # 正常情况
            ind_ok = res['ind_cap_ok']
            raft_ok = res['raft_cap_ok']
            ind_unsat_str = '、'.join(set(res['ind_unsatisfy']))
            raft_unsat_str = '、'.join(set(res['raft_unsatisfy']))

            if ind_ok and raft_ok:
                summary_text += f"{building}采用独立基础或筏形基础时，基底分布地层满足承载力要求，从持力层承载力角度考虑可初步考虑采用天然地基。\n"
            elif not ind_ok and raft_ok:
                summary_text += f"{building}采用独立基础时，基底分布的{ind_unsat_str}不满足承载力要求，从持力层承载力角度考虑不宜采用天然地基；采用筏形基础时，基底分布地层满足承载力要求，从持力层承载力角度考虑可初步考虑采用天然地基。\n"
            elif ind_ok and not raft_ok:  # 理论上较少见，但也写上
                summary_text += f"{building}采用独立基础时，基底分布地层满足承载力要求；采用筏形基础时，基底分布的{raft_unsat_str}不满足承载力要求。\n"
            else:
                # 都不满足
                all_unsat = set(res['ind_unsatisfy'] + res['raft_unsatisfy'])
                summary_text += f"{building}基底分布的{'、'.join(all_unsat)}不满足承载力要求，从持力层承载力角度考虑不宜采用天然地基。\n"

        doc.add_paragraph(summary_text)
        doc.add_paragraph(
            '本次地基承载力修正为估算值，实际计算地基承载力特征值是否满足基底荷载要求时，设计应根据建筑物的基础实际尺寸、荷载大小和分布特征进行详细修正。')

        # 处理软弱下卧层
        # === 筛选需要进行软弱下卧层验算的楼栋 ===
        # 修改逻辑：只要独立基础 OR 筏形基础其中一个满足承载力，就进行验算
        # 1. 筛选需要验算的楼栋（只要有一种基础形式承载力满足，就验算下卧层）
        # 必须定义 weak_buildings 变量，因为后面的代码（第248行）需要用到它
        # 1. 收集所有需要验算的数据（一次性收集）
        weak_buildings_data = {}  # 存储 {楼栋名: 下卧层列表}

        # 筛选条件：只要独立或筏板有一个满足承载力，就进行验算
        check_list = [b for b in buildings_list if
                      building_results[b]['ind_cap_ok'] or building_results[b]['raft_cap_ok']]

        for b in check_list:
            layers = get_weak_underlayers(b, buildings, holes, hole_strata, layer_info)
            if layers:
                weak_buildings_data[b] = layers

        # 2. 生成表格（只执行一次）
        if weak_buildings_data:
            doc.add_paragraph('\n                                                          软弱下卧层承载力验算    表5')
            table5 = doc.add_table(rows=1, cols=6, style='Table Grid')
            h = table5.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text, h[4].text, h[5].text = (
                '拟建建筑', '软弱下卧层', '基础形式', 'pz+pcz(kPa)', 'faz(kPa)', '是否满足')

            weak_ok_list = []  # 两种基础都满足
            weak_partial_list = []  # 只有一种满足
            weak_nok_list = []  # 都不满足

            for building, weak_layers in weak_buildings_data.items():
                B, L, pk = buildings[building].get('width'), buildings[building].get('length'), buildings[building].get(
                    'load')
                if not all((B, L, pk)): continue
                pc = gamma * 1.5

                # 临时标志位
                this_ind_weak_ok = True
                this_raft_weak_ok = True

                for layer in weak_layers:
                    # 计算参数
                    theta_rad = math.radians(calculate_theta(layer['Z'] / B, layer['es1_es2']))
                    l_name = layer['layer_name']
                    fak_w = layer['fak']
                    Z = layer['Z']
                    _, eta_d, _ = get_eta_params(l_name, silt_params, clay_params)

                    # 计算 faz
                    faz = fak_w + eta_d * gamma * (1.5 + Z - 0.5)
                    faz_val = round(faz, 2) if isinstance(faz, (int, float)) else 0
                    faz_str = str(faz_val) if faz_val else '/'

                    # --- 独立基础验算 ---
                    pz_ind = (9 * (pk - pc) / ((3 + 2 * layer['Z'] * math.tan(theta_rad)) ** 2)) + (
                                1.5 + layer['Z']) * gamma
                    ind_sat = (faz_val >= pz_ind)
                    if not ind_sat: this_ind_weak_ok = False

                    r1 = table5.add_row().cells
                    r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text = (
                        building, layer['layer_str'], '独立基础', f"{pz_ind:.2f}", faz_str)
                    r1[5].text = '满足' if ind_sat else '不满足'

                    # --- 筏形基础验算 ---
                    pz_raft = (B * L * (pk - pc) / ((B + 2 * layer['Z'] * math.tan(theta_rad)) * (
                            L + 2 * layer['Z'] * math.tan(theta_rad)))) + (1.5 + layer['Z']) * gamma
                    raft_sat = (faz_val >= pz_raft)
                    if not raft_sat: this_raft_weak_ok = False

                    r2 = table5.add_row().cells
                    r2[0].text, r2[2].text, r2[3].text, r2[4].text = '', '筏形基础', f"{pz_raft:.2f}", faz_str
                    r2[5].text = '满足' if raft_sat else '不满足'

                # === 关键：将结果存回全局 building_results ===
                building_results[building]['ind_weak_ok'] = this_ind_weak_ok
                building_results[building]['raft_weak_ok'] = this_raft_weak_ok

                # 统计用于写总结
                if this_ind_weak_ok and this_raft_weak_ok:
                    weak_ok_list.append(building)
                elif not this_ind_weak_ok and not this_raft_weak_ok:
                    weak_nok_list.append(building)
                else:
                    weak_partial_list.append(building)

            # 3. 生成表5总结
            weak_summary = "根据表5验算结果："
            if weak_ok_list:
                weak_summary += f"{'、'.join(weak_ok_list)}软弱下卧层承载力满足规范要求；"
            if weak_nok_list:
                weak_summary += f"{'、'.join(weak_nok_list)}软弱下卧层承载力不满足规范要求；"
            if weak_partial_list:
                desc_parts = []
                for b in weak_partial_list:
                    i_ok = building_results[b]['ind_weak_ok']
                    if i_ok:
                        desc_parts.append(f"{b}独立基础满足要求但筏形基础不满足")
                    else:
                        desc_parts.append(f"{b}独立基础不满足要求但筏形基础满足")
                weak_summary += "，".join(desc_parts)

            if weak_summary.endswith("；"): weak_summary = weak_summary[:-1]
            doc.add_paragraph(weak_summary + "。")
        else:
            # 如果没有任何楼栋需要验算，最好给个提示或日志
            logging.info("无建筑物需要进行软弱下卧层验算。")

        # ==================== 【表5 逻辑结束】 ====================
                # ==================== 【替换结束】 ====================
            # =================================================

        add_deformation_section(doc, buildings, buildings_list, holes, hole_strata, layer_info, building_results)

        # --- 新增部分：第三节到第五节内容 ---

        # 确定是否有卵石、岩层、黏土、含卵石粉质黏土
        # ==================== 【由 add_deformation_section 之后开始替换】 ====================

        # --- 1. 定义后续章节所需的地层参数 ---
        has_pebble = any('卵石' in info.get('name', '') for info in layer_info.values())
        has_rock = any('岩' in info.get('name', '') for info in layer_info.values())
        has_clay = any('黏土' in info.get('name', '') for info in layer_info.values())
        has_pebble_clay = any('含卵石粉质黏土' in info.get('name', '') for info in layer_info.values())
        has_fill = any(is_fill(lid) for lid in layer_info.keys())

        # 确定复合地基处理方法
        composite_method = "高压旋喷法、CFG桩法" if has_pebble else "CFG桩"

        # 确定复合地基桩端持力层
        if any('中等风化' in info.get('name', '') for info in layer_info.values()):
            composite_hold = "中等风化岩层"
        elif any('强风化' in info.get('name', '') for info in layer_info.values()):
            composite_hold = "强风化岩层"
        else:
            composite_hold = "卵石"

        # 确定桩端持力层字符串
        pile_end_hold_layers = [f"{info['name']}（{lid}）" for lid, info in layer_info.items() if
                                '中等风化' in info.get('name', '')]
        pile_end_hold_str = '、'.join(pile_end_hold_layers) if pile_end_hold_layers else '中密卵石'

        # 辅助函数：格式化地层标签
        def format_layer_label(raw):
            if not raw:
                return ''
            for lid, info in layer_info.items():
                plain = f"{info.get('name', '')}{lid}"
                if raw.startswith(plain):
                    suffix = raw[len(plain):]
                    full_name = f"{info.get('name', '未知层')}（{lid}）"
                    return f"{full_name}{suffix}"
            return raw

        # 初始化 building_hold_layer 字典
        building_hold_layer = {}
        for building in buildings_list:
            base_layers = building_results[building].get('base_layers', [])
            if base_layers:
                most_common = Counter(base_layers).most_common(1)
                building_hold_layer[building] = most_common[0][0] if most_common else ''
            else:
                building_hold_layer[building] = ''

        # --- 2. 重新计算天然地基可行性（生成 第(3)节 内容） ---

        # 显式初始化列表，防止 NameError
        natural_ok_buildings = []  # 最终判定为天然地基可行的楼栋
        natural_not_ok_buildings = []  # 最终判定为不可行的楼栋
        natural_ok_texts = []  # 用于生成第(3)节的文字描述

        # 辅助函数：获取持力层名称
        def get_hold_layer_name(b_name):
            bl = building_results[b_name].get('base_layers', [])
            if bl:
                return '、'.join(sorted(list(set([x.split('（')[0] for x in bl]))))
            return "相应持力层"

        for building in buildings_list:
            res = building_results[building]

            # === 核心判断逻辑 ===
            # 1. 基础承载力、下卧层、变形是否满足
            final_ind_ok = res['ind_cap_ok'] and res['ind_weak_ok'] and res.get('deform_satisfy', True)
            final_raft_ok = res['raft_cap_ok'] and res['raft_weak_ok'] and res.get('deform_satisfy', True)

            # === 2. 特殊修正：如果基底高于地面（全在填土或架空），直接判定为不宜采用天然地基 ===
            if res.get('all_above_ground', False):
                final_ind_ok = False
                final_raft_ok = False

            layer_name = get_hold_layer_name(building)

            if final_ind_ok and final_raft_ok:
                natural_ok_texts.append(
                    f"{building}可考虑采用天然地基，基础形式可考虑独立基础或筏形基础，以{layer_name}作为基础持力层")
                natural_ok_buildings.append(building)
            elif final_ind_ok and not final_raft_ok:
                natural_ok_texts.append(
                    f"{building}可考虑采用天然地基，基础形式可考虑独立基础，以{layer_name}作为基础持力层")
                natural_ok_buildings.append(building)
            elif not final_ind_ok and final_raft_ok:
                natural_ok_texts.append(
                    f"{building}可考虑采用天然地基，基础形式可考虑筏形基础，以{layer_name}作为基础持力层")
                natural_ok_buildings.append(building)
            else:
                # 都不满足
                # 修改：无论是否高于地面，只要不满足上述条件，都加入“不宜采用”列表
                # 这样“基底高于地面”的楼栋就会正确出现在“不宜直接采用天然地基”的结论中
                natural_not_ok_buildings.append(building)

        # --- 3. 输出 第(3)节 文档内容 ---
        doc.add_heading('（3）天然地基方案分析', level=3)
        doc.add_paragraph('根据前述分析（包括地基均匀性、持力层承载力、软弱下卧层验算及地基变形验算）：')

        if natural_ok_texts:
            doc.add_paragraph("；\n".join(natural_ok_texts) + "。")

        if natural_not_ok_buildings:
            doc.add_paragraph(
                f"{'、'.join(natural_not_ok_buildings)}从承载力、软弱下卧层或变形角度考虑，不宜直接采用天然地基。")

        # ==================== 【替换结束，紧接着是 '3、地基处理评价'】 ====================
        # ===============================================

        # 3、地基处理评价
        if natural_not_ok_buildings:
            sheet_param = wb['成都地区地层参数']
            param_dict = {}
            for row in sheet_param.iter_rows(min_row=5, values_only=True):
                if row[0]:
                    name = str(row[0]).strip().replace(' ', '').replace('\u3000', '')
                    state = str(row[1] or '/').strip().replace(' ', '').replace('\u3000', '')
                    key = (name, state)
                    param_dict[key] = {
                        'cfg_qsia': row[3] if row[3] is not None else '/',
                        'cfg_qpa': row[4] if row[4] is not None else '/',
                        'jet_qsia': row[5] if row[5] is not None else '/',
                        'jet_qpa': row[6] if row[6] is not None else '/'
                    }
                    logging.info(f"Added param key {repr(key)}: cfg_qsia={row[3]}, cfg_qpa={row[4]}")  # 日志查看每个key的值
            doc.add_heading(
                '3、地基处理评价',
                level=2)

            # （1）必要性等
            doc.add_heading('（1）地基处理的必要性、处理方法、范围及适宜性', level=3)
            doc.add_paragraph(
                '本项目部分拟建建筑基底分布地层不满足设计要求，在技术可靠、经济的前提下可考虑采用地基处理的方式，如换填、降低基础标高、复合基地等处理措施来有效加固地基，使其承载力及变形满足设计的要求。针对本场地部分拟建建筑基底下地层分布情况，换填、降低基础标高、复合地基处理方式适宜本场地地质条件，可根据设计需要进行地基加固处理。')

            # 计算处理范围建筑
            processing_buildings = []
            for building in buildings_list:

                # === 【新增关键代码】 如果该楼栋已经判定为天然地基OK，则直接跳过 ===
                if building in natural_ok_buildings:
                    continue
                # ============================================================
                building_holes = [h for h, info in holes.items() if building in info.get('building', [])]
                fill_thicknesses = []
                for h in building_holes:
                    ground_elev = holes[h]['elev']
                    base_elev = buildings[building]['embed_elev']
                    depth_base = ground_elev - base_elev
                    strata = hole_strata[h]
                    current_d = 0
                    fill_th = 0
                    started = False
                    for l_id, bottom_d in strata:
                        if bottom_d is None or bottom_d == '':
                            continue
                        layer_th = bottom_d - current_d
                        if current_d >= depth_base:
                            started = True
                        if started and is_fill(l_id):
                            fill_th += layer_th
                        elif started and not is_fill(l_id) and fill_th > 0:
                            break
                        current_d = bottom_d
                    fill_thicknesses.append(fill_th)
                th_lt1 = len([th for th in fill_thicknesses if th < 1])
                th_1to2 = len([th for th in fill_thicknesses if 1 <= th <= 2])
                th_2to3 = len([th for th in fill_thicknesses if 2 < th <= 3])
                if th_lt1 > 0 or (th_1to2 <= 2 and th_2to3 <= 1):
                    processing_buildings.append(building)

            proc_str = '、'.join(processing_buildings) if processing_buildings else ''
            unsatisfy_layers = set()
            for building in processing_buildings:
                unsatisfy_layers.update(
                    building_results[building].get('ind_unsatisfy', []) + building_results[building].get(
                        'raft_unsatisfy', []))
                # === 【修改代码】过滤掉填土层 ===
                # 过滤规则：剔除名称中包含“填土”的层（因为填土已在前半句建议换填）
                # 同时也剔除 ID 以 "1-" 开头的层（双重保险）
                final_unsatisfy_layers = []
                for layer_str in sorted(list(unsatisfy_layers)):
                    # 如果名字里有"填土" 或者 包含 "1-" (填土层号特征)，则跳过
                    if "填土" in layer_str or "1-" in layer_str:
                        continue
                    final_unsatisfy_layers.append(layer_str)

                # 生成最终字符串，如果过滤后为空，保留原代码的默认值 '黏土(2-1)' 以防报错，但通常会有其他层
                unsatisfy_str = '、'.join(final_unsatisfy_layers) if final_unsatisfy_layers else '黏土(2-1)'
                # ==============================

            doc.add_paragraph(
                f"本项目可考虑进行地基处理的范围包括{proc_str}地段。考虑到基底分布的填土厚度不大，可根据设计需要考虑对基底分布的填土进行换填处理或降低基础标高，以及在满足设计要求的前提下对{unsatisfy_str}等不满足要求的地层进行复合地基处理，可考虑{composite_method}复合地基。处理后的地基稳定性整体较好。")

            # （2）方案
            doc.add_heading('（2）地基处理方案', level=3)
            no_fill_not_ok = [b for b in natural_not_ok_buildings if b not in processing_buildings]
            no_fill_str = '、'.join(no_fill_not_ok) if no_fill_not_ok else ''
            base_layer_no_fill = '、'.join(set([l for b in no_fill_not_ok for l in
                                               building_results.get(b, {}).get('base_layers',
                                                                               [])])) if no_fill_not_ok else '黏土(2-1)'
            doc.add_paragraph(
                f"{no_fill_str}基底主要分布为{base_layer_no_fill}，可考虑采用{composite_method}复合地基进行地基处理，复合地基可考虑以{composite_hold}作为CFG桩桩端持力层。" if no_fill_str else '')
            fill_layer_str = '填土(1-1、1-2)' if has_fill else ''
            not_satisfy_proc = unsatisfy_str
            doc.add_paragraph(
                f"{proc_str}基底分布的{not_satisfy_proc}不满足要求，考虑到基底分布的{fill_layer_str}厚度不大，可考虑将基底厚度不大的{fill_layer_str}进行混凝土换填处理，或换填碎石并夯实挤密后采用{composite_method}复合地基进行地基处理，复合地基可考虑以{composite_hold}作为CFG桩桩端持力层。" if proc_str else '')

            # （3）设计参数
            doc.add_heading('（3）复合地基设计参数建议', level=3)
            doc.add_paragraph(
                '复合地基方案应委托具备相应资质单位进行专项设计，复合地基承载力应通过具备相应资质的单位进行试验检测确定。正式施工成桩前，应选取具有代表性的地段进行试桩，试桩满足设计要求后方可进行大面积施工。有关复合地基的设计参数建议值见下表。根据设计需要，桩端阻力特征值参数也可考虑通过试验检测进行确定。')
            # ==================== 【替换开始：表7 和 表8 生成逻辑（剔除填土）】 ====================
            doc.add_paragraph('                                              有关CFG桩法复合地基的相关设计指标建议值表      表7')

            # 重新读取参数表（确保逻辑连贯）
            sheet_param = wb['成都地区地层参数']
            param_dict = {}
            for row in sheet_param.iter_rows(min_row=5, values_only=True):
                if row[3]:
                    # 加强清理
                    name = str(row[3]).replace(' ', '').replace('\u3000', '').strip()
                    state = str(row[4] or '/').replace(' ', '').replace('\u3000', '').strip()
                    key = (name, state)
                    param_dict[key] = {
                        'cfg_qsia': row[6] if row[6] is not None else '/',
                        'cfg_qpa': row[7] if row[7] is not None else '/',
                        'jet_qsia': row[8] if row[8] is not None else '/',
                        'jet_qpa': row[9] if row[9] is not None else '/',
                        'name': row[3] if row[3] is not None else '/',
                        'state': row[4] if row[4] is not None else '/',
                        'dry_qsik': row[10] if row[10] is not None else '/',
                        'dry_qpk': row[11] if row[11] is not None else '/',
                        'mud_qsik': row[12] if row[12] is not None else '/',
                        'mud_qpk': row[13] if row[13] is not None else '/',
                        'pre_qsik': row[14] if row[14] is not None else '/',
                        'pre_qpk': row[15] if row[15] is not None else '/',
                        'rc_natural': row[16] if row[16] is not None else '/',
                        'rc_saturated': row[17] if row[17] is not None else '/',
                        'm_coeff': row[18] if row[18] is not None else '/',
                        'k_coeff': row[19] if row[19] is not None else '/'
                    }

            table7 = doc.add_table(rows=1, cols=3, style='Table Grid')
            hdr7 = table7.rows[0].cells
            hdr7[0].text = '岩土名称'
            hdr7[1].text = '桩侧阻力特征值qsia(kPa)'
            hdr7[2].text = '桩端阻力特征值qpa(kPa)'

            # 遍历所有地层
            for layer_id in sorted(layer_info.keys()):
                info = layer_info[layer_id]
                name = info['name']
                state = info['state']

                # === 修改点：严格过滤填土 ===
                # 排除条件：ID是填土格式(1-xx) OR 名称包含"填土"
                if is_fill(layer_id) or '填土' in name:
                    continue
                # ==========================

                full_name = f"{name}（{layer_id}）"
                logging.info(f"Table7 - Processing layer {layer_id}: name={repr(name)}, state={repr(state)}")

                key = (name, state) if ('黏土' in name or '粉质黏土' in name) else (name, '/')

                if key in param_dict:
                    qsia = param_dict[key]['cfg_qsia']
                    qpa = param_dict[key]['cfg_qpa']
                else:
                    # 备用查找
                    fallback_key = (name, '/')
                    if fallback_key in param_dict:
                        qsia = param_dict[fallback_key]['cfg_qsia']
                        qpa = param_dict[fallback_key]['cfg_qpa']
                    else:
                        qsia = '/'
                        qpa = '/'

                r = table7.add_row().cells
                r[0].text = full_name
                r[1].text = str(qsia)
                r[2].text = str(qpa)

            if has_pebble:
                doc.add_paragraph('                                            有关高压旋喷桩法复合地基的相关设计指标建议值表      表8')
                table8 = doc.add_table(rows=1, cols=3, style='Table Grid')
                hdr8 = table8.rows[0].cells
                hdr8[0].text = '岩土名称'
                hdr8[1].text = '桩侧阻力特征值qsia(kPa)'
                hdr8[2].text = '桩端阻力特征值qpa(kPa)'

                for layer_id in sorted(layer_info.keys()):
                    info = layer_info[layer_id]
                    name = info['name']
                    state = info['state']

                    # === 修改点：严格过滤填土 ===
                    if is_fill(layer_id) or '填土' in name:
                        continue
                    # ==========================

                    full_name = f"{name}（{layer_id}）"

                    key = (name, state) if ('黏土' in name or '粉质黏土' in name) else (name, '/')

                    if key in param_dict:
                        qsia = param_dict[key]['jet_qsia']
                        qpa = param_dict[key]['jet_qpa']
                    else:
                        no_state_key = (name, '/')
                        if no_state_key in param_dict:
                            qsia = param_dict[no_state_key]['jet_qsia']
                            qpa = param_dict[no_state_key]['jet_qpa']
                        else:
                            qsia = '/'
                            qpa = '/'

                    r = table8.add_row().cells
                    r[0].text = full_name
                    r[1].text = str(qsia)
                    r[2].text = str(qpa)
            # ==================== 【替换结束】 ====================
            # （4）风险
            doc.add_heading('（4）地基处理设计施工可能遇到的风险及对环境的影响', level=3)
            doc.add_paragraph(
                '根据场地工程地质条件，地基处理设计方案不合理、不完善或与实际情况不符，可能导致基础处理承载力不足、施工过程中出现问题以及导致后续使用中出现安全隐患等问题。施工过程中可能出现因软弱土、松散岩土体、膨胀土及地下水体等因素导致的坑壁或孔壁坍塌、施工材料不适宜、质量不达标或对相邻地基造成扰动影响。地基处理过程中可能产生大量的废水、废渣等污染物，施工过程中可能会破坏周围的植被、土壤等生态环境，如果地基处理不当可能会对地下水、地表水等自然资源造成污染。')

            # （5）注意和检测
            doc.add_heading('（5）地基处理应注意的问题和检测建议', level=3)
            doc.add_paragraph(
                '场地场平后场地内整体较平坦，具备地基处理施工机具要求，部分地段必要时可考虑路面硬化。场地内暂未通水通电。施工前应详细调场地内及周边地下管网的分布情况，并对存在影响的管网进行迁改。场地分布的填土为欠固结土，欠固结土不宜采用复合地基方式进行地基处理，宜对欠固结土先采用碎石换填并夯实挤密方式进行预处理。对场地周围存在地表水下渗或地下管网渗漏的导致岩土体性状变差的应采取加固措施。')
            if has_clay and has_fill:
                doc.add_paragraph(
                    '场地分布的填土夹有碎石或块石，黏土中夹有分化岩屑，且黏土具有胀缩性，可能对复合地基成桩施工造成影响，适合机械旋挖（钻）成孔，成孔过程中可能会发生塌孔现象，必要时施工过程中可采取有效的措施控制施工质量，可采取套筒、泥浆护壁钻进引孔措施。')
            if has_clay:
                doc.add_paragraph(
                    '场地硬塑黏土厚度较大或风化基岩区域由于静压或锤击法桩基不宜穿透，故黏土厚度较大或风化基岩区域不适宜采用挤土桩施工。')
            if has_rock:
                doc.add_paragraph('场地基岩面存在起伏，复合地基需注意桩间土的协同受力及差异沉降问题。')
            doc.add_paragraph(
                '地基处理应进行专项岩土设计，同时应注意设计方案及施工工艺的合理适宜性，选择合适的施工材料，地下水富集时水位以下可能需要采用水下混凝土施工工艺，严格把控施工质量以满足设计要求并及时验收合格，同时注意减小对桩间土的扰动及对环境的影响。')
            doc.add_paragraph(
                '若采用复合地基，成桩结束应待桩身强度达到设计要求后，由具有法定资格的检测单位进行低应变检测、钻芯检取样检查桩身完整性及桩底沉渣情况，采用载荷试验等手段进行复合地基承载力及单桩承载力检测及桩间土承载力验证等工作。正式施工成桩前，应选取具有代表性的地段进行试桩，试桩满足设计要求后方可进行大面积施工。若采用碎石等进行换填处理，必要时建议对换填处理后碎石的密实度进行检测。')

        # 4、桩基础评价
        if natural_not_ok_buildings:
            doc.add_heading(
                '4、桩基础评价',
                level=2)
            pile_buildings_str = '、'.join(natural_not_ok_buildings)
            pile_unsatisfy_layer = unsatisfy_str
            doc.add_paragraph(
                f"{pile_buildings_str}基底分布{pile_unsatisfy_layer}地段由于基底分布地层不满足要求，也可考虑采用桩基础。")

            doc.add_heading('（1）成桩可行性分析', level=3)
            doc.add_paragraph(
                '目前常用的桩基类型主要有：干作业法钻孔灌注桩、泥浆护壁钻孔灌注桩及PHC预应力管桩。桩型选择应充分考虑拟建物结构特性、场区地层条件、周边环境条件及同类工程经验。')
            doc.add_paragraph('根据本场地工程地质条件以及当地同类工程经验，各桩型的可行性分析如下：')
            doc.add_paragraph('a.干作业法钻孔灌注桩')
            doc.add_paragraph(
                '干作业法钻孔灌注桩适用于地下水贫乏且不易出现孔壁塌落的场地，成孔工艺有多种，如机械钻孔法、旋挖钻进法。结合本项目实际情况，本场地部分地段存在厚度不等的填土，极易产生缩颈或孔壁塌落，工艺适宜性稍差，若要考虑干作业法钻孔灌注桩，局部地段可能需结合采取护筒等措施，采用旋挖成孔较为适宜。' if has_fill else '干作业法钻孔灌注桩适用于地下水贫乏且不易出现孔壁塌落的场地，成孔工艺有多种，如机械钻孔法、旋挖钻进法。结合本项目实际情况，采用旋挖成孔较为适宜。')
            doc.add_paragraph('b.泥浆护壁钻孔灌注桩')
            doc.add_paragraph(
                '泥浆护壁钻孔灌注桩适用性较广，对易出现孔壁塌落的场地效果较好，成孔工艺采用机械成孔，可在地下水位以下施工，成孔简单，工艺可行。结合本项目实际情况，采用旋挖成孔的泥浆护壁钻孔灌注桩较为适宜。')
            doc.add_paragraph('c. PHC预应力管桩')
            # 收集所有中间风险部分
            middle_parts = []

            if has_fill:
                middle_parts.append('本场地的填土中夹碎石或块石，')

            if has_rock:
                middle_parts.append('PHC预应力管桩难以穿进中等风化岩层，')

            # 拼接成完整的一段文字
            full_text = (
                    'PHC预应力管桩适用于松散软弱地层为主或岩性较差的场地，成桩工艺有静压法、锤击法。结合本项目实际情况，'
                    + ''.join(middle_parts).rstrip('，')  # 去掉多余的逗号
                    + '，采用该桩型最后可能形成的结果是无法达到预期设计的桩端持力层及桩长，易出现断桩的情况，且成桩施工将对地基扰动过大，故PHC预应力管桩的适宜性应根据现场试桩结果确定。'
            )

            # 如果有中间内容，确保逗号衔接自然；如果没有，直接连接
            if middle_parts:
                # 已经在上面处理了 rstrip('，')，这里可以再加一个判断避免空逗号
                full_text = full_text.replace('，采用该桩型', '，采用该桩型')  # 正常情况
            else:
                full_text = full_text.replace('，采用该桩型', '，采用该桩型')  # 无中间时保持原样

            doc.add_paragraph(full_text)
            if has_pebble:
                doc.add_paragraph(
                    '卵石层中零星分布有粒径较大的漂石，在施工过程中可能会出现管桩不易穿透漂石富集区域，甚至无法穿越甚至无法到达设计持力层，采用该桩型最后形成的结果是因侧阻力较低导致桩长无法预计，且成桩施工将对地基土扰动过大。故建议进行现场试桩，以确定预应力管桩对该工程的适用性。')

            doc.add_heading('（2）桩基础型式的对比、施工条件、地下水的影响及桩基施工对环境的影响', level=3)
            doc.add_paragraph(
                '桩基础型式的对比：干作业法钻孔灌注桩主要适用于地下水贫乏且不易出现孔壁塌落的场地，成孔工艺有多种，如机械钻孔法、旋挖钻进法。' + (
                    '针对本场地地质的特殊性，本场地存在厚度不等的填土，极易产生孔壁塌落，工艺适宜性较差，若要考虑干作业法钻孔灌注桩，需结合采取护筒等措施，采用旋挖成孔最为适宜。' if has_fill else '') + '其成孔易于控制，桩底沉渣易于清除，对桩间土和桩端土的扰动小，有利于桩间土和桩端土工程力学性能的最大发挥，但可能的局部地下水富集地段对干作业法影响较大，不利于桩基施工质量控制，将对环境产生一定噪音及振动影响。泥浆护壁钻孔灌注桩适用性较广，对易出现孔壁塌落的场地效果相对较好，成孔工艺采用机械成孔，可在地下水位以下施工，成孔简单，能一定程度上降低施工过程中的噪音及振动影响，但桩底沉渣不易完全清除，护壁泥浆可能对桩间土和桩端土产生浸泡扰动，一定程度上不利于桩间土和桩端土工程力学性能的发挥。' + (
                    'PHC预应力管桩难以穿进中等风化岩层，采用该桩型最后可能形成的结果是无法达到预期设计的桩端持力层及桩长，易出现断桩的情况，且成桩施工将对地基扰动过大，故PHC预应力管桩的适宜性应根据现场试桩结果确定。' if has_rock else '') + (
                    '卵石层中零星分布有粒径较大的漂石，在施工过程中可能会出现管桩不易穿透漂石富集区域，甚至无法穿越甚至无法到达设计持力层，采用该桩型最后形成的结果是因侧阻力较低导致桩长无法预计，且成桩施工将对地基土扰动过大。故建议进行现场试桩，以确定预应力管桩对该工程的适用性。' if has_pebble else '') + '结合本项目实际情况，综合分析，采用旋挖成孔的泥浆护壁钻孔灌注桩较为适宜。')
            doc.add_paragraph(
                '施工条件：场平后进行桩基施工时，场地整体平坦，具备大型器械施工作业条件，局部区域可能需要进行硬化处理，以防止器械侧翻出现安全事故。')
            doc.add_paragraph(
                '地下水的影响：桩基施工时，地下水的流动可能导致孔壁塌落，对施工产生安全隐患，同时产生较多沉渣，必要时可采取适宜的护壁措施；地下水及施工用水对岩土浸泡也可能降低桩身及桩端土层物理力学性质，降低单桩承载力，孔内地下水影响桩体混凝土的灌注，若灌注过程不能完全挤出孔内地下水，可能导致桩身强度降低甚至断桩情况的发生，灌注前桩底沉渣应尽可能清底干净，保障桩底沉渣厚度不大于50mm；桩身混凝土浇筑后，采用“后注浆”工艺，可以一定程度上弥补因水和机械的扰动对桩侧、桩端土体结构强度消弱，并减少桩底沉渣对成桩质量的影响。')
            doc.add_paragraph(
                '桩基施工对环境的影响：桩基施工过程中对地下水体产生扰动影响，泥浆或施工用水可能导致局部地下水体遭受一定程度的污染。在施工过程中会产生大量的泥浆，基坑内进行泥浆的排放困难很大，而且施工中产生的泥浆、施工用水会浸泡基坑底部地层，致使其物理力学性质变差，同时其桩底会产生较多的沉渣，该沉渣难以处理，给桩基础的承载力和变形带来较大的不利影响，同时在施工过程中对环境产生较大的噪音影响。桩基施工将对环境产生一定影响。')

            doc.add_heading('（3）桩端持力层建议及桩型规格、桩长分析', level=3)
            doc.add_paragraph(
                f"若拟建建筑采用桩基础，建议以{pile_end_hold_str}作为桩端持力层。根据本项目场地地层分布特征，结合区域地质及工程经验，若采用连续稳定中等风化岩作为桩端持力层，勘探深度范围内桩端以下一般无洞穴、临空面、破碎岩体或软弱岩层等，必要时亦可通过一桩一孔施工勘察进行探明。" if pile_end_hold_layers else f"若拟建建筑采用桩基础，建议以{pile_end_hold_str}作为桩端持力层。")

            def parse_optional_float(val):
                try:
                    return float(val) if val is not None else None
                except (ValueError, TypeError):
                    return None

            # 以下代码移出 def 外面，无缩进
            buildings_list = []
            building_to_embed_elev = {}  # 新增：建筑到基底标高的字典
            sheet_buildings = wb['1.1']
            for row in sheet_buildings.iter_rows(min_row=6, values_only=True):  # 移除max_row，自动到末尾
                if row and len(row) > 0 and row[2]:  # 修改回row[0]，检查建筑名称列
                    building_name = str(row[2]).strip()
                    if building_name and building_name != '拟建建筑名称':  # 排除表头，如果误读
                        buildings_list.append(building_name)
                        embed_elev = parse_optional_float(row[11])  # 基础埋深标高在row[11]
                        if embed_elev is not None:
                            building_to_embed_elev[building_name] = embed_elev

            # Parse layers from sheet "2.4各孔地层"
            layers = defaultdict(list)
            current_hole = None
            sheet_layers = wb['2.4各孔地层']
            for row in sheet_layers.iter_rows(min_row=2, values_only=True):
                if row:
                    hole = row[0]
                    if hole:
                        current_hole = str(hole).strip()
                    layer_id = str(row[1]).strip() if row[1] else None
                    bottom_depth = parse_optional_float(row[2])
                    if layer_id and bottom_depth is not None:
                        layers[current_hole].append((layer_id, bottom_depth))

            # Define hold layer sets
            medium_weathered_rock = {'4-2', '5-3', '5-4'}  # 中等风化岩
            dense_pebble = set()  # 密实卵石 (if layer codes available, add them)
            mid_dense_pebble = {'2-2'} if len(medium_weathered_rock) == 0 else set()  # 中密卵石 fallback, adjust as needed
            hold_sets = [medium_weathered_rock, dense_pebble, mid_dense_pebble]

            # Function to get top depth of the first matching hold layer
            def get_hold_top_depth(layers, hole, hold_sets):
                layer_list = layers.get(hole, [])
                prev_bottom = 0.0
                for layer_id, bottom in layer_list:
                    for hold_set in hold_sets:
                        if layer_id in hold_set:
                            return prev_bottom
                    prev_bottom = bottom
                return None

            holes = {}
            sheet_holes = wb['1.5单孔']
            for row in sheet_holes.iter_rows(min_row=2, values_only=True):
                if row and row[0]:  # 确保有孔号
                    hole_id = str(row[0]).strip()
                    elev = parse_optional_float(row[1])
                    buildings_str = row[12] if len(row) > 12 else ''  # 安全检查长度
                    buildings = [b.strip() for b in re.split(r'[、,]', buildings_str) if
                                 b.strip()] if buildings_str else []
                    if elev is not None:
                        holes[hole_id] = {'elev': elev, 'buildings': buildings}  # 始终设置'buildings'，即使空

            # Compute min top elev for each building
            building_hold_elevs = defaultdict(list)
            for hole_id, info in holes.items():
                top_depth = get_hold_top_depth(layers, hole_id, hold_sets)
                if top_depth is not None:
                    top_elev = info['elev'] - top_depth
                    for b in info['buildings']:
                        building_hold_elevs[b].append(top_elev)

            # Create table9
            doc.add_paragraph('根据地区工程经验，对可能采用的桩型规格和桩端高程建议如下表9。')
            doc.add_paragraph('                                                           桩型规格和桩端高程建议表      表9')
            table9 = doc.add_table(rows=1, cols=4, style='Table Grid')
            hdr9 = table9.rows[0].cells
            hdr9[0].text = '拟建建筑'
            hdr9[1].text = '可采用的桩型'
            hdr9[2].text = '常用规格范围(mm)'
            hdr9[3].text = '桩端高程建议(m)'
            pile_types = ['干作业法钻孔灌注桩', '泥浆护壁钻孔灌注桩', 'PHC预应力管桩']
            specs = ['Ф800～Ф1400', 'Ф600～Ф1400', 'Ф300～Ф600']

            # 1. 动态定义有效的桩端持力层ID集合
            valid_hold_ids = set()
            for lid, info in layer_info.items():
                name = info.get('name', '')
                # 关键词匹配：中等风化、密实/中密/稍密卵石
                if any(kw in name for kw in ['中等风化', '中风化', '密实卵石', '中密卵石', '稍密卵石', '卵石']):
                    # 排除含卵石粉质黏土、填土、松散层
                    if '含卵石' not in name and '填土' not in name and '松散' not in name:
                        valid_hold_ids.add(lid)
            logging.info(f"识别出的有效桩端持力层ID: {valid_hold_ids}")

            for building in buildings_list:
                # === 【新增筛选逻辑】 ===
                # 如果该楼栋已被判定为天然地基可行，则不在表9中列出
                if building in natural_ok_buildings:
                    continue
                # =======================

                # 获取该建筑的基础埋深标高
                embed_elev = building_to_embed_elev.get(building)

                # 收集该建筑下所有钻孔计算出的【建议标高1：地层顶-0.5】
                candidates_limit_1 = []

                # 获取该建筑关联的钻孔
                assoc_holes = [h for h, info in holes.items() if building in info.get('buildings', [])]

                if embed_elev is not None and assoc_holes:
                    for hole_id in assoc_holes:
                        hole_elev = holes[hole_id].get('elev')
                        strata = hole_strata.get(hole_id, [])

                        if hole_elev is None or not strata:
                            continue

                        # 计算基底在钻孔中的深度 (相对于孔口)
                        base_depth_in_hole = max(0.0, hole_elev - embed_elev)

                        # 寻找“基底以下”的第一个持力层
                        found_top_depth = None
                        prev_bottom = 0.0

                        for lid, bottom in strata:
                            effective_bottom = bottom if bottom is not None else 999.0

                            # 逻辑：如果这一层的底部 比 基底深度 深，说明这一层在基底以下（或包含基底）
                            if effective_bottom > base_depth_in_hole:
                                if lid in valid_hold_ids:
                                    # 找到了符合条件的持力层，取其顶深
                                    found_top_depth = prev_bottom
                                    break

                            prev_bottom = effective_bottom

                        if found_top_depth is not None:
                            # 换算为绝对标高
                            layer_top_elev = hole_elev - found_top_depth
                            candidates_limit_1.append(layer_top_elev)

                # 计算最终建议标高
                final_suggest_elev = '/'

                if embed_elev is not None:
                    # 限制条件2：基底标高 - 6.0m
                    limit_2 = embed_elev - 6.0

                    limit_1 = None
                    if candidates_limit_1:
                        # 规则：各钻孔取最大值 (最浅的那个持力层顶板)
                        max_layer_top = max(candidates_limit_1)
                        # 规则：减去 0.5m
                        limit_1 = max_layer_top - 0.5

                    if limit_1 is not None:
                        # 规则：取小值 (min)
                        val = min(limit_1, limit_2)
                        final_suggest_elev = f"{val:.1f}"
                    else:
                        # 兜底：如果没有探测到持力层，按最小桩长6m控制
                        final_suggest_elev = f"{limit_2:.1f}"

                # 填充表格
                for i in range(3):
                    row = table9.add_row().cells
                    if i == 0:
                        row[0].text = building
                    else:
                        row[0].text = ''  # 合并单元格效果

                    row[1].text = pile_types[i]
                    row[2].text = specs[i]

                    if final_suggest_elev != '/':
                        row[3].text = f'不高于{final_suggest_elev}'
                    else:
                        row[3].text = '不高于/'

                        # ==================== 表9生成结束 ====================

            doc.add_heading('（4）桩基设计参数建议及单桩竖向承载力的确定', level=3)
            doc.add_paragraph(
                '本工程灌注桩单桩承载力应通过现场桩静载试验确定，也可采用深层平板载荷试验确定桩端土层的极限端阻力或承载力特征值。试验前可按表10建议参数设计。')
            doc.add_paragraph('                                                               桩基设计参数建议值表      表10')
            sheet_param = wb['成都地区地层参数']
            param_dict = {}
            for row in sheet_param.iter_rows(min_row=5, values_only=True):
                if row[3]:
                    # 加强清理：移除所有空格和全角空格
                    name = str(row[3]).replace(' ', '').replace('\u3000', '').strip()
                    state = str(row[4] or '/').replace(' ', '').replace('\u3000', '').strip()
                    key = (name, state)
                    param_dict[key] = {
                        'cfg_qsia': row[6] if row[6] is not None else '/',
                        'cfg_qpa': row[7] if row[7] is not None else '/',
                        'jet_qsia': row[8] if row[8] is not None else '/',
                        'jet_qpa': row[9] if row[9] is not None else '/',
                        'name': row[3] if row[3] is not None else '/',
                        'state': row[4] if row[4] is not None else '/',
                        'dry_qsik': row[10] if row[10] is not None else '/',
                        'dry_qpk': row[11] if row[11] is not None else '/',
                        'mud_qsik': row[12] if row[12] is not None else '/',
                        'mud_qpk': row[13] if row[13] is not None else '/',
                        'pre_qsik': row[14] if row[14] is not None else '/',
                        'pre_qpk': row[15] if row[15] is not None else '/',
                        'rc_natural': row[17] if row[17] is not None else '/',
                        'rc_saturated': row[18] if row[18] is not None else '/',
                        'm_coeff': row[19] if row[19] is not None else '/',
                        'k_coeff': row[20] if row[20] is not None else '/'
                    }
                    # 假设has_clay, has_fill, has_rock, has_pebble等已定义，例如：
                    has_clay = any('黏土' in info['name'] for info in layer_info.values())
                    has_fill = any(is_fill(lid) for lid in layer_info)
                    has_rock = any('岩' in info['name'] for info in layer_info.values())
                    has_pebble = any('卵石' in info['name'] for info in layer_info.values())
                    # has_pebble_clay 等类似，根据需要定义
            table10 = doc.add_table(rows=1, cols=11, style='Table Grid')
            # === 【新增】设置列宽 ===
            table10.autofit = False  # 关键：关闭自动适应，否则宽度设置无效

            # 设置第一列宽度为 2cm
            table10.columns[0].width = Cm(1.82)
            table10.columns[1].width = Cm(1.6)
            table10.columns[2].width = Cm(1.6)
            table10.columns[3].width = Cm(1.6)
            table10.columns[4].width = Cm(1.6)
            table10.columns[5].width = Cm(1.6)
            table10.columns[6].width = Cm(1.6)
            table10.columns[7].width = Cm(1.32)
            table10.columns[8].width = Cm(1.32)
            table10.columns[9].width = Cm(1.6)
            table10.columns[10].width = Cm(1.6)
            # =======================
            hdr10 = table10.rows[0].cells
            hdr10[0].text = '岩土名称'
            hdr10[1].text = '干作业法挖(钻)孔灌注桩桩侧阻力特征值qsia(kPa)'
            hdr10[2].text = '干作业法挖(钻)孔灌注桩桩端阻力特征值qpa(kPa)'
            hdr10[3].text = '泥浆护壁钻(冲)孔灌注桩桩侧阻力特征值qsia(kPa)'
            hdr10[4].text = '泥浆护壁钻(冲)孔灌注桩桩端阻力特征值qpa(kPa)'
            hdr10[5].text = '预应力管桩桩侧阻力特征值qsia(kPa)'
            hdr10[6].text = '预应力管桩桩端阻力特征值qpa(kPa)'
            hdr10[7].text = '天然单轴抗压强度(MPa)'
            hdr10[8].text = '饱和天然单轴抗压强度(MPa)'
            hdr10[9].text = '灌注桩地基土水平抗力系数的比例系数m(MN/m4)'
            hdr10[10].text = '灌注桩岩体水平抗力系数k(MN/ m4)'
            # 只添加非填土地层
            for layer_id in sorted(layer_info.keys()):
                info = layer_info[layer_id]
                name = info['name']
                state = info['state']
                full_name = f"{name}（{layer_id}）"
                logging.info(f"Table7 - Processing layer {layer_id}: name={repr(name)}, state={repr(state)}")

                key = (name, state) if ('黏土' in name or '粉质黏土' in name) else (name, '/')
                logging.info(f"Table7 - Trying key: {repr(key)}")

                if key in param_dict:
                    qsik1 = param_dict[key]['dry_qsik']
                    qpk2 = param_dict[key]['dry_qpk']
                    qsik3 = param_dict[key]['mud_qsik']
                    qpk4 = param_dict[key]['mud_qpk']
                    qsik5 = param_dict[key]['pre_qsik']
                    qpk6 = param_dict[key]['pre_qpk']
                    rcn = param_dict[key]['rc_natural']
                    rcs = param_dict[key]['rc_saturated']
                    mc = param_dict[key]['m_coeff']
                    kc = param_dict[key]['k_coeff']
                else:
                    qsik1 = '/'
                    qpk2 = '/'
                    qsik3 = '/'
                    qpk4 = '/'
                    qsik5 = '/'
                    qpk6 = '/'
                    rcn = '/'
                    rcs = '/'
                    mc = '/'
                    kc = '/'

                r = table10.add_row().cells
                r[0].text = full_name
                r[1].text = str(qsik1)
                r[2].text = str(qpk2)
                r[3].text = str(qsik3)
                r[4].text = str(qpk4)
                r[5].text = str(qsik5)
                r[6].text = str(qpk6)
                r[7].text = str(rcn)
                r[8].text = str(rcs)
                r[9].text = str(mc)
                r[10].text = str(kc)
                fill_neg = '、'.join([f"填土（{lid}）" for lid in layer_info if is_fill(lid)])
            doc.add_paragraph(f"注：{fill_neg}负摩阻力系数建议取0.3。")
            if has_rock:
                doc.add_paragraph(
                    '由于岩石抗压试验与岩石作为地基时的实际受力状态存在明显差异，大直径灌注桩单桩承载力的现场试验结果往往要比按岩石抗压强度估算的结果高很多，因此从发挥地基潜力的角度讲，大直径灌注桩的单桩承载力也应根据现场试验确定。')

            doc.add_heading('（5）桩基础沉降特征', level=3)
            settle_hold = '中等风化岩' if pile_end_hold_layers else ('中密卵石或密实卵石' if has_pebble else '中密卵石')
            doc.add_paragraph(
                f"若桩端进入{settle_hold}中，如采用适宜的桩基方案，桩基沉降根据工程经验一般能满足设计要求，必要时可进行专项咨询。")

            doc.add_heading('（6）特殊性岩土及地下水对桩基危害及措施', level=3)
            if has_fill:
                doc.add_paragraph(
                    '根据场平设计，场地部分地段存在填土，基坑后期也将进行回填。由于填土自身的自重固结尚未完成，在桩基形成后，该部分土体对桩身可能产生负摩阻效应。基桩设计应考虑该层的负摩阻效应。未经按设计压实处理的场地后期经场平形成的填土（即场平扰动后的或回填后的填土）建议负摩阻力系数可参考0.30取值，天然状态下综合内摩察角建议取值20°，饱和状态下综合内摩察角建议取值15°。建议设计或施工时应采取相应措施如对填土进行预压、增设保护桩、安装套筒或在桩身涂滑动薄膜（如涂沥青）等降低负摩阻力，以及采取措施如采取加大桩基截面尺寸、加强桩基配筋、设置抗滑桩等方法消除或降低斜坡上回填土对桩基础水平剪力的不利作用。场地分布的填方地基，为防止后期环境地面过度下沉，建议填方分层压实，压实系数按设计及相关规范执行。')
            if has_clay:
                doc.add_paragraph(
                    '膨胀土对桩会产生胀缩力，同时采用锤击法进行预制桩施工时“挤土效应”十分明显，亦会造成断桩、斜桩现象，设计时应充分考虑膨胀土膨胀而对桩产生的向上拔力，对于挤土桩必要时可考虑引孔等措施。桩基成孔时由于胀缩可能导致孔壁失稳垮塌，必要时可考虑套筒措施确保桩基施工正常进行。')
            if has_rock:
                doc.add_paragraph(
                    '场地分布有易风化岩，桩基成孔后应及时清底并浇筑混凝土，避免风化岩长时间浸水软化或暴露风化导致桩周及桩端地层工程性质变差。成孔过程应控制钻进速度，避免过渡扰动。')
            if has_pebble_clay:
                doc.add_paragraph(
                    '场地分布的含卵石粉质黏土属混合土。含卵石粉质黏土呈黄褐、灰褐等色，物质组成主要为粉质黏土及卵石，来源为地层形成过程中场地所在区域该部分物质沉积而成，粉质黏土呈可塑状，卵石含量一般不大于35%，粒径多在2～8cm，局部夹少量砾石或大于20cm的块石，颗粒呈强～中等风化状态，均匀性较差，水平向变化差异较小，垂直向卵石颗粒含量略有增加趋势。局部卵石富集，其自稳定性较差，桩基施工时，应注意扰动作用下引起的垮孔、沉渣等问题，可考虑采用泥浆护壁或套筒措施。考虑到混合土的不均匀性，若以其作为基础持力层，应剔除其中的粗大颗粒。')
            doc.add_paragraph('地下水对特殊性岩土的影响：')
            if has_fill:
                doc.add_paragraph(
                    '填土被地下水浸泡扰动后性状变差且强度大幅下降，开挖易出现垮塌或掉块现象，易产生变形及不均匀沉降，对相关埋设的管线、地坪、建筑物也会有一定的沉降影响。')
            if has_clay:
                doc.add_paragraph(
                    '膨胀土具有遇水快速膨胀、失水快速收缩的特点，在地下水作用下易出现含水率变化而产生胀缩变形破坏，同时地下水浸泡会导致浸泡影响区强度大幅下降，承载力大大降低。')
            if has_pebble_clay:
                doc.add_paragraph(
                    '混合土在地下水作用下孔隙结构易发生改变，导致土体结构松散，影响混合土的工程性状，降低土体强度，可能导致土体失稳。')
            if has_rock:
                doc.add_paragraph(
                    '风化岩在地下水作用下易出现风化加剧及发生软化，导致地基承载力的衰减、抗剪切性指标的降低，对其稳定性也有一定影响。')
            doc.add_paragraph(
                '建议施工前先降低地下水位，采取做好疏排水、硬化隔水、铺设保护膜或保护层等措施，避免对土体过渡扰动，严禁大量长时间的浸水暴晒，以降低或消除地下水的不良影响。')

            doc.add_heading('（7）成桩可能遇到的风险', level=3)
            doc.add_paragraph(
                '场地部分地段分布的松散岩土体自身稳定性较差，成桩施工过程中可能出现孔壁塌落、缩颈，造成成桩困难且对施工安全及桩基工程质量产生影响。桩底沉渣未清除或沉渣厚度较大，可能导致桩端作用无法发挥到预期效果，进而影响桩基工程质量。成孔时护壁用的泥浆流失过快，可能影响成孔质量，对成桩施工安全性及桩基工程质量产生影响。桩基施工过程中产生的振动可能对临近基础工程质量及周围构筑物产生影响，产生的噪音及泥浆材料等可能对环境造成影响，施工材料及用水也可能对地下水体造成影响。护壁泥浆用量异常，可能对桩基工程质量产生一定影响。局部地下水富集地段可能出现可能孔壁塌落，同时水下沉渣不易清除，也可能通过浸泡导致桩间土及桩端土理力学性质降低，影响桩基工程质量。孔内地下水影响桩体混凝土的灌注，若灌注过程不能完全挤出孔内地下水，可能导致桩身强度降低甚至断桩情况的发生。桩基施工作业时，场平后表层松散土体可能形成不均匀沉降导致器械侧翻出现安全事故，必要时地面需要进行硬化处理。')

            doc.add_heading('（8）桩基设计施工应注意的问题和检测建议', level=3)
            doc.add_paragraph(
                '桩基设计应充分考虑场地土的稳定性和承载力，选择合适的桩基平面布置、桩型和桩长。施工前制定合理的桩基施工方案，并采取相应的措施保证施工安全。成孔时应注意护壁泥浆用量，避免流失过快影响成孔质量及环境。混凝土灌注前桩底沉渣应尽可能清底干净，保障桩底沉渣厚度不大于50mm。桩身混凝土浇筑后，可考虑采用“后注浆”工艺，可以一定程度上弥补因水和机械的扰动对桩侧、桩端土体结构强度消弱，并减少桩底沉渣对成桩质量的影响。桩基施工过程中应注意振动及噪音控制，避免对临近基础工程质量及周围环境产生影响。施工材料及用水应注意环保问题，避免对环境及地下水体造成影响。桩底沉渣应及时清除，避免影响桩端作用发挥效果。在软弱土地段施工时，应采取相应的措施如泥浆护壁或刚护筒确保顺利成孔并保证施工安全和桩基工程质量。若采用挤土桩，可能出现桩端无法穿进持力层的情况，必要时可考虑引孔措施。必要时可通过桩基施工勘察探明桩端以下一定深度范围内是否存在洞穴、临空面、破碎岩体或软弱岩层等。桩基混凝土浇筑前，应及时对桩端持力层进行验收，避免长时间浸泡或暴露导致桩端土桩间土及桩端土理力学性质降低，影响桩基工程质量。')
            doc.add_paragraph(
                '对桩基础建议根据《建筑地基检测技术规范》（JGJ 340-2015）规定进行相关检测。建议通过载荷试验等方式进行单桩承载力检测，并对桩身完整性及桩底沉渣厚度等进行检测。')

        # 5、地基基础方案及持力层建议
        # ================================
        # 表11：拟建物基础方案及持力层建议一览表（终极稳定版 + 已修复 building_info）
        # ================================
        doc.add_heading('5、地基基础方案及持力层建议', level=2)
        doc.add_paragraph('综合分析，拟建建筑地基基础方案建议见表11。')
        doc.add_paragraph('                                                  拟建物基础方案及持力层建议一览表 表11')
        table11 = doc.add_table(rows=1, cols=2, style='Table Grid')
        # === 【新增】设置列宽 ===
        table11.autofit = False  # 关键：关闭自动适应，否则宽度设置无效

        # 设置第一列宽度为 2cm
        table11.columns[0].width = Cm(2.5)

        # 建议：设置第二列宽度（A4纸除去页边距约剩16cm，这里设为14cm填满剩余空间）
        # 如果不设置第二列，Word有时会显示异常，建议显式指定
        table11.columns[1].width = Cm(14)
        # =======================
        hdr11 = table11.rows[0].cells
        hdr11[0].text = '建筑物名称'
        hdr11[1].text = '基础形式与持力层建议'

        # ==================== 1. 读取建筑信息（必须放在最前面！）===================
        building_info = {}
        try:
            sheet_1_1 = wb['1.1']
            for row in sheet_1_1.iter_rows(min_row=6, values_only=True):
                if not row or len(row) < 3 or not row[2]:
                    continue
                name = str(row[2]).strip()
                if name in {'拟建建筑名称', ''}:
                    continue
                embed_elev = parse_optional_float(row[11]) if len(row) > 11 else None
                load = parse_optional_float(row[12]) if len(row) > 12 else 0
                building_info[name] = {'embed_elev': embed_elev, 'load': load}
        except Exception as e:
            logging.error(f"读取1.1表建筑信息失败: {e}")
            messagebox.showerror("错误", f"读取1.1表建筑信息失败: {e}")

        # ==================== 2. 辅助函数 ====================
        def get_building_holes(building):
            return [h for h, info in holes.items() if building in [str(b).strip() for b in info.get('buildings', [])]]

        def get_embed_depth(hole_id, building):
            elev = holes.get(hole_id, {}).get('elev')
            embed_elev = building_info.get(building, {}).get('embed_elev')
            if embed_elev is None:
                embed_elev = buildings.get(building, {}).get('embed_elev')
            if elev is None or embed_elev is None: return 0.0
            return max(elev - embed_elev, 0)

        def safe_layers(hole_id):
            raw = hole_strata.get(hole_id, [])
            if not isinstance(raw, list): return []
            return [(str(lid).strip(), float(bot)) for lid, bot in raw if lid and bot is not None]

        def get_layer_full_name(lid):
            info = layer_info.get(lid, {})
            return f"{info.get('name', '未知层')}（{lid}）"

        fill_layers = {'1', '1-1', '1-2'}

        # ==================== 3. 天然地基持力层 ====================
        def get_natural_hold_layer(building):
            """
            【真实岩土版】天然地基持力层
            规则：把这栋楼所有钻孔在基底标高处实际所在的地层全部列出来
            如果基底在填土层，则找填土层下第一个非填土层
            用"、"连接，填土不显示
            """
            actual_layers = set()  # 用 set 去重
            debug = []

            for h in get_building_holes(building):
                depth = get_embed_depth(h, building)
                layer_seq = safe_layers(h)
                debug.append(f"{h}: 层序 {layer_seq}")
                if depth <= 0.01:
                    debug.append(f"{h}: 基底高于地面")
                    continue

                acc = 0.0
                found = None
                at_fill = False  # 标记基底是否在填土层

                for lid, bot in layer_seq:
                    if acc + (bot - acc) >= depth:  # 基底落在这层
                        if lid not in fill_layers:
                            found = lid
                            break
                        else:
                            # 基底在填土层，继续往下找第一个非填土层
                            at_fill = True
                    elif at_fill:
                        # 已经过了基底填土层，找到第一个非填土层
                        if lid not in fill_layers:
                            found = lid
                            break
                    acc = bot

                # 兜底：如果还没找到，从最后一层往前找
                if found is None and layer_seq:
                    for lid, _ in reversed(layer_seq):
                        if lid not in fill_layers:
                            found = lid
                            break

                if found:
                    name = get_layer_full_name(found)
                    actual_layers.add(name)
                    debug.append(f"{h}: {depth:.2f}m → {name}")
                else:
                    debug.append(f"{h}: 未找到非填土地层")

            # 打印调试（运行时看日志就知道每个孔落在哪）
            logging.info(f"\n{'=' * 50}")
            logging.info(f"{building} 天然地基持力层真实分布")
            for line in debug:
                logging.info(f"  {line}")
            logging.info(f"  → 最终写入报告：{','.join(actual_layers) or '无'}")
            logging.info(f"{'=' * 50}\n")

            if not actual_layers:
                return "未知层"

            # 用"、"连接，多个层就全写上
            return "、".join(sorted(actual_layers))

        # ==================== 4. 桩端持力层（仅四类）===================
        def get_pile_hold_layer(building):
            candidates = []
            for h in get_building_holes(building):
                for lid, bot in reversed(safe_layers(h)):
                    if lid in fill_layers: continue
                    name = layer_info.get(lid, {}).get('name', '')
                    if any(kw in name for kw in ['中等风化', '中风化']):
                        candidates.append(lid)
                        break
                    if '密实卵石' in name:
                        candidates.append(lid)
                        break
                    if '中密卵石' in name:
                        candidates.append(lid)
                        break
                    if '稍密卵石' in name:
                        candidates.append(lid)
                        break
                    break  # 其他层直接跳过

            if not candidates:
                return None

            from collections import Counter
            count = Counter(candidates)

            # 优先级1：所有中等风化岩
            medium = [lid for lid in count if
                      any(kw in layer_info.get(lid, {}).get('name', '') for kw in ['中等风化', '中风化'])]
            if medium:
                return '、'.join(get_layer_full_name(lid) for lid in medium)

            # 优先级2-4
            for keyword in ['密实卵石', '中密卵石', '稍密卵石']:
                matches = [lid for lid in count if keyword in layer_info.get(lid, {}).get('name', '')]
                if matches:
                    best = max(matches, key=count.get)
                    return get_layer_full_name(best)

            return None

        # ==================== 5. 其他辅助函数 ====================
        def has_thick_pebble(building):
            total, cnt = 0, 0
            for h in get_building_holes(building):
                depth = get_embed_depth(h, building)
                cur = thick = 0.0
                for lid, bot in safe_layers(h):
                    if '卵石' in layer_info.get(lid, {}).get('name', ''):
                        thick += max(0, min(bot, depth + 100) - cur)
                    cur = bot
                    if cur >= depth: break
                if thick > 0:
                    total += thick
                    cnt += 1
            return cnt > 0 and total / cnt >= 6.0

        def get_unsatisfied_layers(building):
            lids = set()
            for h in get_building_holes(building):
                depth = get_embed_depth(h, building)
                cur = 0.0
                for lid, bot in safe_layers(h):
                    if cur >= depth: break
                    if bot >= depth and layer_info.get(lid, {}).get('bearing_capacity', 0) < 150:
                        lids.add(lid)
                    cur = bot
            return [get_layer_full_name(lid) for lid in lids]

        def is_processable(thick_list):
            if not thick_list: return False
            if max(thick_list) < 1: return True
            c1_2 = sum(1 for t in thick_list if 1 <= t <= 2)
            c2_3 = sum(1 for t in thick_list if 2 < t <= 3)
            return c1_2 <= 2 and c2_3 <= 1

        def get_thickness_list(building, condition):
            thicks = []
            for h in get_building_holes(building):
                depth = get_embed_depth(h, building)
                cur = 0.0
                for lid, bot in safe_layers(h):
                    if cur >= depth: break
                    if condition(lid):
                        thicks.append(round(min(bot, depth) - cur, 2))
                    cur = bot
            return [t for t in thicks if t > 0]

        def is_below_fill_all_satisfied(building):
            for h in get_building_holes(building):
                depth = get_embed_depth(h, building)
                cur = 0.0
                passed_fill = False
                for lid, bot in safe_layers(h):
                    if cur >= depth: break
                    if lid in fill_layers:
                        passed_fill = True
                    elif passed_fill:
                        if layer_info.get(lid, {}).get('bearing_capacity', 0) < 150:
                            return False
                    cur = bot
            return True

        # ==================== 6. 主循环 ====================
        def is_processable(thick_list):
            if not thick_list:
                return False

            # 1. 检查是否有任何孔厚度超过 3m -> 直接不满足
            # 对应条件：基底以下...厚度＞3m的钻孔的个数≤0个
            if any(t > 3.0 for t in thick_list):
                return False

            # 2. 统计特定区间的孔数
            # 厚度≥1m且≤2m
            c1_2 = sum(1 for t in thick_list if 1.0 <= t <= 2.0)
            # 厚度＞2m且≤3m
            c2_3 = sum(1 for t in thick_list if 2.0 < t <= 3.0)

            # 3. 判断数量限制
            # 条件：1-2m个数≤2个 并且 2-3m个数≤1个
            # (小于1m的孔不影响此判断，因为它们既不违反>3m，也不占用计数名额)
            return c1_2 <= 2 and c2_3 <= 1

        # 辅助函数：获取基底以下【连续填土】的厚度列表和名称集合
        def get_continuous_fill_info(b_name):
            thicks = []
            names = set()
            for h in get_building_holes(b_name):
                embed_depth = get_embed_depth(h, b_name)
                cur_top = 0.0
                total_th = 0.0
                hole_names = set()

                for lid, bot in safe_layers(h):
                    # 跳过基底之上的层
                    if bot <= embed_depth:
                        cur_top = bot
                        continue

                    start_d = max(cur_top, embed_depth)

                    if lid in fill_layers:
                        total_th += (bot - start_d)
                        hole_names.add(get_layer_full_name(lid))
                    else:
                        break  # 连续填土结束

                    cur_top = bot

                if total_th > 0:
                    thicks.append(round(total_th, 2))
                    names.update(hole_names)
            return thicks, sorted(list(names))

        # 辅助函数：获取基底以下【连续不满足承载力地层】的厚度列表和名称集合
        def get_continuous_weak_info(b_name, building_load):
            thicks = []
            names = set()
            for h in get_building_holes(b_name):
                embed_depth = get_embed_depth(h, b_name)
                cur_top = 0.0
                total_th = 0.0
                hole_names = set()

                for lid, bot in safe_layers(h):
                    if bot <= embed_depth:
                        cur_top = bot
                        continue

                    start_d = max(cur_top, embed_depth)
                    fak = layer_info.get(lid, {}).get('bearing_capacity', 0)

                    # 判定标准：承载力 < 楼栋荷载
                    # 只要不满足，就累加厚度（中砂+松散卵石会在这里合并）
                    if fak < building_load:
                        total_th += (bot - start_d)
                        hole_names.add(get_layer_full_name(lid))
                    else:
                        break  # 遇到满足承载力的层，连续弱层结束

                    cur_top = bot

                if total_th > 0:
                    thicks.append(round(total_th, 2))
                    names.update(hole_names)
            return thicks, sorted(list(names))

        # 辅助函数：检查“待处理层”之下的地层是否满足要求
        def check_below_block_ok(b_name, is_block_layer_func, building_load):
            all_ok = True
            for h in get_building_holes(b_name):
                embed_depth = get_embed_depth(h, b_name)
                passed_block = False
                hole_ok = True

                for lid, bot in safe_layers(h):
                    if bot <= embed_depth: continue

                    if is_block_layer_func(lid):
                        passed_block = True
                    elif passed_block:
                        # 穿过待处理层后，检查紧邻的一层
                        fak = layer_info.get(lid, {}).get('bearing_capacity', 0)
                        if fak < building_load:
                            hole_ok = False
                        break  # 只检查紧邻的一层

                if not hole_ok:
                    all_ok = False
                    break
            return all_ok

        # 辅助函数：计算基底以下所有“卵石”类地层的总厚度（用于旋喷桩判定）
        def get_total_pebble_thickness(b_name):
            avg_th_list = []
            for h in get_building_holes(b_name):
                embed_depth = get_embed_depth(h, b_name)
                cur_top = 0.0
                th = 0.0
                for lid, bot in safe_layers(h):
                    if bot <= embed_depth:
                        cur_top = bot
                        continue
                    start_d = max(cur_top, embed_depth)
                    name = layer_info.get(lid, {}).get('name', '')
                    if '卵石' in name:
                        th += (bot - start_d)
                    cur_top = bot
                avg_th_list.append(th)
            if not avg_th_list: return 0
            return sum(avg_th_list) / len(avg_th_list)

        for building in buildings_list:
            row = table11.add_row().cells
            row[0].text = building

            load = building_info.get(building, {}).get('load')
            if load in (None, 0):
                load = buildings.get(building, {}).get('load', 0)
            load = load or 0

            # 获取各维度状态
            res = building_results[building]

            # === 再次执行核心判断 ===
            final_ind_ok = res['ind_cap_ok'] and res['ind_weak_ok'] and res.get('deform_satisfy', True)
            final_raft_ok = res['raft_cap_ok'] and res['raft_weak_ok'] and res.get('deform_satisfy', True)

            natural_layer = get_natural_hold_layer(building)
            pile_layer = get_pile_hold_layer(building)
            hold_layer_str = pile_layer if pile_layer else '相应持力层'

            # 辅助：旋喷桩判断
            pebble_th = get_total_pebble_thickness(building)
            jet_ok = (load <= 650 and pebble_th >= 6.0)

            # --- 分支1：天然地基可行（任意一种形式）---
            if final_ind_ok or final_raft_ok:
                # 获取持力层名称（优先用前面统计的，如果没有则用计算的）
                h_raw = building_hold_layer.get(building, '')
                hold_layer = format_layer_label(h_raw) or natural_layer

                if final_ind_ok and final_raft_ok:
                    txt = f"建议采用天然地基，基础形式可考虑独立基础或筏形基础，以{hold_layer}作为基础持力层。"
                elif final_ind_ok:
                    txt = f"建议采用天然地基，基础形式可考虑独立基础，以{hold_layer}作为基础持力层。"
                elif final_raft_ok:
                    txt = f"建议采用天然地基，基础形式可考虑筏形基础，以{hold_layer}作为基础持力层。"

            # --- 分支2：地下室（且天然地基不可行）---
            elif '地下室' in building:
                unsat_names = get_unsatisfied_layers(building)  # 当前基底下不满足承载力的层
                unsat_str = '、'.join(unsat_names) if unsat_names else '无'

                # 新增：计算“穿过不满足层后”的真实持力层（即紧邻其下的满足承载力的层）
                good_layers_after_unsat = set()
                for h in get_building_holes(building):
                    depth = get_embed_depth(h, building)
                    layer_seq = safe_layers(h)
                    passed_unsat = False
                    for lid, bot in layer_seq:
                        if bot <= depth:
                            continue
                        fak = layer_info.get(lid, {}).get('bearing_capacity', 0)
                        load_here = building_info.get(building, {}).get('load', 0) or buildings.get(building, {}).get(
                            'load', 0) or 0

                        if not passed_unsat:
                            if fak < load_here:
                                passed_unsat = True  # 正在穿过不满足层
                                continue
                            else:
                                # 当前就在满足的层上，不需要降低标高
                                good_layers_after_unsat.add(get_layer_full_name(lid))
                                break
                        else:
                            # 已经穿过不满足层，当前层就是目标持力层
                            good_layers_after_unsat.add(get_layer_full_name(lid))
                            break

                good_layer_str = '、'.join(sorted(good_layers_after_unsat)) if good_layers_after_unsat else '满足承载力要求的地层'

                txt = (
                    f"建议基底未分布{unsat_str}地段可考虑采用天然地基；\n"
                    f"基底分布有{unsat_str}地段建议采用独立基础或筏板，对{unsat_str}进行混凝土换填，以换填后的人工地基作为基础持力层，"
                    f"或降低基底标高，穿过{unsat_str}，以{good_layer_str}作为基础持力层，"
                    f"或采用桩基础，以{hold_layer_str}作为桩端持力层。"
                )

            # --- 分支3：其他不满足的情况（填土、承载力不足等）---
            else:
                # 判断基底以下是否有填土
                fill_thicks, fill_names_list = get_continuous_fill_info(building)
                has_fill_below = bool(fill_thicks)

                if has_fill_below:
                    # 有填土
                    fill_name = '、'.join(fill_names_list) or '填土'
                    jet_suffix = "也可换填填土后采用高压旋喷桩复合地基，以处理后的复合地基作为持力层。" if jet_ok else ""

                    if is_processable(fill_thicks):
                        # 填土不厚，建议换填
                        txt = (
                            f"可采用对{fill_name}换填，基础形式可考虑独立基础或筏形基础，"
                            f"或换填填土后采用CFG桩法复合地基，复合地基以{hold_layer_str}作为CFG桩桩端持力层，"
                            f"或采用桩基础，以{hold_layer_str}作为桩端持力层。"
                            f"{jet_suffix}"
                        )
                    else:
                        # 填土太厚
                        txt = f"建议采用桩基础，以{hold_layer_str}作为桩端持力层。"

                else:
                    # 无填土，纯粹是承载力或下卧层不够
                    weak_thicks, weak_names_list = get_continuous_weak_info(building, load)
                    weak_str = '、'.join(weak_names_list) or '不满足承载力地层'
                    jet_suffix = "也可采用高压旋喷桩复合地基，以处理后的复合地基作为持力层。" if jet_ok else ""

                    if is_processable(weak_thicks):
                        txt = (
                            f"可采用对{weak_str}换填，基础形式可考虑独立基础或筏形基础，"
                            f"或采用CFG桩法复合地基，复合地基以{hold_layer_str}作为CFG桩桩端持力层，"
                            f"或采用桩基础，以{hold_layer_str}作为桩端持力层。"
                            f"{jet_suffix}"
                        )
                    else:
                        txt = (
                            f"可采用CFG桩法复合地基，复合地基以{hold_layer_str}作为CFG桩桩端持力层，"
                            f"或采用桩基础，以{hold_layer_str}作为桩端持力层。"
                            f"{jet_suffix}"
                        )

            row[1].text = txt

        doc.add_paragraph('注：基础持力层的选择应以结构设计验算是否满足荷载和变形要求为准。')

        # 为所有表格应用字体（包括其他表格，如果有）
        for table in doc.tables:
            set_table_font(table)



        doc.save(output_path)
        wb.close()


    except FileNotFoundError as e:
        logging.error(f"文件未找到: {e}")
        raise
    except KeyError as e:
        logging.error(f"Excel中缺少工作表或列: {e}")
        raise
    except Exception as e:
        logging.error(f"未知错误: {e}", exc_info=True)
        raise


def run_analysis_with_ui(template_path, input_path, output_path):
    try:
        run_analysis(template_path, input_path, output_path, messagebox.askyesno)
        messagebox.showinfo("完成", "分析完成，Word 报告已生成。")
    except FileNotFoundError as e:
        messagebox.showerror("错误", f"文件未找到: {e}")
    except KeyError as e:
        messagebox.showerror("错误", f"Excel中缺少工作表或列: {e}")
    except Exception as e:
        messagebox.showerror("错误", f"发生未知错误: {e}")


# --- GUI 界面 (现代化UI) ---
class ModernUI:
    # 颜色主题
    PRIMARY_BLUE = "#2563EB"      # 主蓝色
    PRIMARY_GREEN = "#16A34A"     # 主绿色
    LIGHT_BLUE = "#EFF6FF"        # 浅蓝背景
    LIGHT_GREEN = "#F0FDF4"       # 浅绿背景
    WHITE = "#FFFFFF"
    GRAY_50 = "#F9FAFB"
    GRAY_100 = "#F3F4F6"
    GRAY_200 = "#E5E7EB"
    GRAY_400 = "#9CA3AF"
    GRAY_600 = "#4B5563"
    GRAY_800 = "#1F2937"
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("成都地区地基基础分析及选型")
        self.root.geometry("420x580")
        self.root.configure(bg=self.WHITE)
        self.root.resizable(False, False)
        
        # 变量
        self.input_path_var = tk.StringVar()
        self.output_path_var = tk.StringVar()
        self.template_path_var = tk.StringVar()
        self.current_step = 1
        
        self.setup_ui()
    
    def setup_ui(self):
        # 主容器
        main_frame = tk.Frame(self.root, bg=self.WHITE, padx=25, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题区域
        self.create_header(main_frame)
        
        # 步骤指示器
        self.create_step_indicator(main_frame)
        
        # 功能卡片区域
        self.create_cards(main_frame)
        
        # 底部署名
        self.create_footer(main_frame)
    
    def create_header(self, parent):
        header_frame = tk.Frame(parent, bg=self.WHITE)
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        # 主标题
        title_label = tk.Label(
            header_frame, 
            text="成都地区地基基础分析及选型",
            font=("Microsoft YaHei UI", 16, "bold"),
            fg=self.GRAY_800,
            bg=self.WHITE
        )
        title_label.pack(anchor='w')
        
        # 副标题
        subtitle_label = tk.Label(
            header_frame,
            text="Foundation Analysis & Selection System",
            font=("Segoe UI", 9),
            fg=self.GRAY_400,
            bg=self.WHITE
        )
        subtitle_label.pack(anchor='w', pady=(2, 0))
        
        # 关闭按钮
        close_btn = tk.Label(
            header_frame,
            text="×",
            font=("Arial", 18),
            fg=self.GRAY_400,
            bg=self.WHITE,
            cursor="hand2"
        )
        close_btn.place(relx=1.0, rely=0, anchor='ne')
        close_btn.bind("<Button-1>", lambda e: self.root.destroy())
        close_btn.bind("<Enter>", lambda e: close_btn.configure(fg=self.GRAY_600))
        close_btn.bind("<Leave>", lambda e: close_btn.configure(fg=self.GRAY_400))
    
    def create_step_indicator(self, parent):
        step_frame = tk.Frame(parent, bg=self.WHITE)
        step_frame.pack(fill=tk.X, pady=(0, 25))
        
        steps = [
            ("01", "上传数据"),
            ("02", "开始分析"),
            ("03", "获取结果")
        ]
        
        self.step_circles = []
        self.step_labels = []
        
        for i, (num, text) in enumerate(steps):
            # 步骤容器
            step_container = tk.Frame(step_frame, bg=self.WHITE)
            step_container.pack(side=tk.LEFT, expand=True)
            
            # 圆形数字
            is_active = (i + 1) == self.current_step
            circle_color = self.PRIMARY_BLUE if is_active else self.WHITE
            text_color = self.WHITE if is_active else self.GRAY_400
            border_color = self.PRIMARY_BLUE if is_active else self.GRAY_200
            
            circle = tk.Frame(step_container, bg=circle_color, width=40, height=40)
            circle.pack()
            circle.pack_propagate(False)
            
            # 边框效果
            circle.configure(highlightbackground=border_color, highlightthickness=2)
            
            num_label = tk.Label(
                circle,
                text=num,
                font=("Segoe UI", 11, "bold"),
                fg=text_color,
                bg=circle_color
            )
            num_label.place(relx=0.5, rely=0.5, anchor='center')
            
            # 步骤文字
            step_text = tk.Label(
                step_container,
                text=text,
                font=("Microsoft YaHei UI", 9),
                fg=self.GRAY_600 if is_active else self.GRAY_400,
                bg=self.WHITE
            )
            step_text.pack(pady=(5, 0))
            
            self.step_circles.append((circle, num_label))
            self.step_labels.append(step_text)
            
            # 连接线（除了最后一个）
            if i < len(steps) - 1:
                line_frame = tk.Frame(step_frame, bg=self.WHITE, width=40)
                line_frame.pack(side=tk.LEFT, fill=tk.Y, pady=20)
                line = tk.Frame(line_frame, bg=self.GRAY_200, height=2)
                line.pack(fill=tk.X, pady=18)
    
    def update_step(self, step):
        self.current_step = step
        for i, ((circle, num_label), step_label) in enumerate(zip(self.step_circles, self.step_labels)):
            is_active = (i + 1) <= step
            is_current = (i + 1) == step
            
            circle_color = self.PRIMARY_BLUE if is_active else self.WHITE
            text_color = self.WHITE if is_active else self.GRAY_400
            border_color = self.PRIMARY_BLUE if is_active else self.GRAY_200
            
            circle.configure(bg=circle_color, highlightbackground=border_color)
            num_label.configure(bg=circle_color, fg=text_color)
            step_label.configure(fg=self.GRAY_600 if is_current else self.GRAY_400)
    
    def create_cards(self, parent):
        cards_frame = tk.Frame(parent, bg=self.WHITE)
        cards_frame.pack(fill=tk.BOTH, expand=True)
        
        # 卡片1: 选择Excel文件 (蓝色)
        self.create_card(
            cards_frame,
            icon="",
            title="选择Excel文件",
            subtitle="点击上传地质勘察数据",
            bg_color=self.PRIMARY_BLUE,
            command=self.select_input_file,
            path_var=self.input_path_var
        )
        
        # 卡片2: 输出文件位置 (灰色边框)
        self.create_card(
            cards_frame,
            icon="",
            title="输出文件位置",
            subtitle="结果将自动下载到默认文件夹",
            bg_color=None,
            command=self.select_output_file,
            path_var=self.output_path_var,
            is_outline=True
        )
        
        # 卡片3: 开始分析 (绿色)
        self.create_card(
            cards_frame,
            icon="",
            title="开始分析",
            subtitle="点击运行地基基础智能分析",
            bg_color=self.PRIMARY_GREEN,
            command=self.run_analysis
        )
    
    def create_card(self, parent, icon, title, subtitle, bg_color, command, path_var=None, is_outline=False):
        # 卡片外框
        if is_outline:
            card = tk.Frame(parent, bg=self.WHITE, highlightbackground=self.GRAY_200, highlightthickness=1)
        else:
            card = tk.Frame(parent, bg=bg_color)
        
        card.pack(fill=tk.X, pady=8, ipady=12)
        card.configure(cursor="hand2")
        
        # 内容容器
        content_bg = self.WHITE if is_outline else bg_color
        content = tk.Frame(card, bg=content_bg, padx=15)
        content.pack(fill=tk.BOTH, expand=True)
        
        # 左侧图标
        icon_frame = tk.Frame(content, bg=content_bg, width=45, height=45)
        icon_frame.pack(side=tk.LEFT, padx=(0, 12))
        icon_frame.pack_propagate(False)
        
        if is_outline:
            icon_frame.configure(highlightbackground=self.GRAY_200, highlightthickness=1)
        
        icon_label = tk.Label(
            icon_frame,
            text=icon,
            font=("Segoe UI Emoji", 16),
            bg=content_bg,
            fg=self.PRIMARY_BLUE if is_outline else self.WHITE
        )
        icon_label.place(relx=0.5, rely=0.5, anchor='center')
        
        # 中间文字
        text_frame = tk.Frame(content, bg=content_bg)
        text_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        title_color = self.GRAY_800 if is_outline else self.WHITE
        subtitle_color = self.GRAY_400 if is_outline else "#E0E7FF"
        
        title_label = tk.Label(
            text_frame,
            text=title,
            font=("Microsoft YaHei UI", 12, "bold"),
            fg=title_color,
            bg=content_bg,
            anchor='w'
        )
        title_label.pack(fill=tk.X)
        
        # 如果有路径变量，显示路径或默认副标题
        if path_var:
            if not hasattr(self, 'subtitle_labels'):
                self.subtitle_labels = {}
            sub_label = tk.Label(
                text_frame,
                text=subtitle,
                font=("Microsoft YaHei UI", 9),
                fg=subtitle_color,
                bg=content_bg,
                anchor='w'
            )
            sub_label.pack(fill=tk.X)
            self.subtitle_labels[str(id(path_var))] = (sub_label, subtitle, subtitle_color, content_bg)
        else:
            tk.Label(
                text_frame,
                text=subtitle,
                font=("Microsoft YaHei UI", 9),
                fg=subtitle_color,
                bg=content_bg,
                anchor='w'
            ).pack(fill=tk.X)
        
        # 右侧箭头
        arrow_label = tk.Label(
            content,
            text=">",
            font=("Arial", 16, "bold"),
            fg=self.GRAY_400 if is_outline else "#E0E7FF",
            bg=content_bg
        )
        arrow_label.pack(side=tk.RIGHT)
        
        # 绑定点击事件
        def on_click(e):
            command()
        
        for widget in [card, content, icon_frame, icon_label, text_frame, title_label, arrow_label]:
            widget.bind("<Button-1>", on_click)
        
        # 悬停效果
        def on_enter(e):
            if not is_outline:
                card.configure(bg=self._lighten_color(bg_color))
                for w in [content, icon_frame, icon_label, text_frame, title_label, arrow_label]:
                    try:
                        w.configure(bg=self._lighten_color(bg_color))
                    except Exception:
                        pass
        
        def on_leave(e):
            if not is_outline:
                card.configure(bg=bg_color)
                for w in [content, icon_frame, icon_label, text_frame, title_label, arrow_label]:
                    try:
                        w.configure(bg=bg_color)
                    except Exception:
                        pass
        
        for widget in [card, content, icon_frame, icon_label, text_frame, title_label, arrow_label]:
            widget.bind("<Enter>", on_enter)
            widget.bind("<Leave>", on_leave)
    
    def _lighten_color(self, hex_color):
        """使颜色变亮"""
        if not hex_color:
            return self.GRAY_100
        hex_color = hex_color.lstrip('#')
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        r = min(255, int(r * 1.1))
        g = min(255, int(g * 1.1))
        b = min(255, int(b * 1.1))
        return f"#{r:02x}{g:02x}{b:02x}"
    
    def create_footer(self, parent):
        footer = tk.Label(
            parent,
            text="中建西勘院 文兴",
            font=("Microsoft YaHei UI", 10),
            fg=self.GRAY_400,
            bg=self.WHITE
        )
        footer.pack(side=tk.BOTTOM, pady=(15, 0))
    
    def select_input_file(self):
        path = filedialog.askopenfilename(
            title="选择数据 Excel 文件", 
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if path:
            self.input_path_var.set(path)
            self.update_step(1)
            key = str(id(self.input_path_var))
            if hasattr(self, 'subtitle_labels') and key in self.subtitle_labels:
                label, _, _, _ = self.subtitle_labels[key]
                filename = os.path.basename(path)
                label.configure(text=f"已选择: {filename}")
    
    def select_output_file(self):
        path = filedialog.asksaveasfilename(
            title="保存输出 Word 文件", 
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")]
        )
        if path:
            self.output_path_var.set(path)
            self.update_step(2)
            key = str(id(self.output_path_var))
            if hasattr(self, 'subtitle_labels') and key in self.subtitle_labels:
                label, _, _, _ = self.subtitle_labels[key]
                filename = os.path.basename(path)
                label.configure(text=f"保存至: {filename}")
    
    def run_analysis(self):
        input_path = self.input_path_var.get()
        output_path = self.output_path_var.get()
        
        if not input_path or not output_path:
            messagebox.showwarning("警告", "请先选择数据Excel文件和输出Word路径")
            return
        
        self.update_step(3)
        run_analysis_with_ui(self.template_path_var.get(), input_path, output_path)
    
    def run(self):
        self.root.mainloop()


# 启动应用
if __name__ == "__main__":
    app = ModernUI()
    app.run()