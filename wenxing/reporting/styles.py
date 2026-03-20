# -*- coding: utf-8 -*-
"""
Word 文档样式设置模块
"""

from docx.shared import Pt, Cm
from docx.oxml.ns import qn


def force_set_all_fonts(doc, font_name='宋体', font_size=Pt(10.5)):
    """强制设置文档中所有文字的字体"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run.font.size = font_size


def set_global_styles(doc, font_name='宋体', font_size=Pt(10.5)):
    """设置全局样式"""
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = font_size
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 设置标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = font_name
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(14)
        elif i == 2:
            heading_style.font.size = Pt(12)
        else:
            heading_style.font.size = Pt(10.5)

    # 设置页边距
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)


def set_table_font(table, font_name='宋体', font_size=Pt(9)):
    """设置表格字体"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = font_size
